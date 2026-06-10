from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/burakcoskun/Mocapexpo")
TEMPLATE = Path("/Users/burakcoskun/Downloads/BSM498_Sablon_20161.docx")
OUT = ROOT / "reports" / "Mocap_Bitirme_Kitapcigi_Taslak.docx"


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def style_name(document: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {style.name for style in document.styles}
    return preferred if preferred in names else fallback


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for run_item in paragraph.runs:
        run_item.font.name = "Times New Roman"
        run_item.font.size = Pt(10)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def add_paragraph(document: Document, text: str, style: str, first: bool = False):
    p = document.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Inches(0.35) if first else None
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def add_paragraphs(document: Document, texts: list[str], style: str, first: bool = True) -> None:
    for text in texts:
        add_paragraph(document, text, style, first=first)


def add_heading(document: Document, text: str, style: str):
    p = document.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    return p


def add_caption(document: Document, text: str, caption_style: str):
    p = document.add_paragraph(style=caption_style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text.startswith("Tablo"):
        p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.italic = True
    return p


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    prevent_row_split(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        docx_row = table.add_row()
        prevent_row_split(docx_row)
        cells = docx_row.cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    document.add_paragraph()
    return table


def add_diagram_box(document: Document, lines: list[str]):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F7FA")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if index < len(lines) - 1:
            run.add_break()
    document.add_paragraph()


def front_table(document: Document, rows: list[tuple[str, str]]):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
    document.add_paragraph()
    return table


def add_intro_expansion(document: Document, body_style: str, caption_style: str) -> None:
    add_paragraphs(
        document,
        [
            "Mocap'ın proje olarak değerlendirilmesinde en önemli noktalardan biri, sistemin yalnızca bir algoritma denemesi değil, gerçek kullanıcı akışı olan bir ürün prototipi olmasıdır. Kullanıcı uygulamayı açtığında önce proje ve çekim bağlamı seçer, ardından çekim moduna göre solo, dual veya pro kayıt süreci başlar. Kayıt tamamlandıktan sonra video yalnızca cihaz belleğinde bırakılmaz; take modeliyle ilişkilendirilir, metadata paketi üretilir, upload aşamasına hazırlanır ve backend kayıtlarıyla eşleştirilir. Bu zincirin her halkası rapor açısından ayrı bir mühendislik kararıdır.",
            "Hareket yakalama problemini mobil cihazla çözmeye çalışırken üç temel zorluk öne çıkar. Birinci zorluk, mobil kameranın sabit stüdyo kamerası gibi davranmamasıdır; elde tutulan cihazda sarsıntı, otomatik odak, otomatik pozlama ve kare zamanı değişimleri görülebilir. İkinci zorluk, insan pozunun videodan çıkarılmasının belirsiz bir problem olmasıdır; tek kamerada derinlik ve yer teması doğrudan ölçülemez. Üçüncü zorluk ise kullanıcı tarafından üretilen video dosyasının büyük ve hassas bir veri olmasıdır; bu dosya güvenli taşınmalı, izlenebilir biçimde işlenmeli ve sonuçları doğrulanabilir olmalıdır.",
            "Bu nedenle proje tasarımında 'çekim yapmak' ile 'hareket yakalamak' birbirinden ayrılmıştır. Çekim yapmak mobil uygulamanın sorumluluğudur; hareket yakalama ise backend worker tarafında yürütülen model, optimizasyon, temizleme ve export süreçlerinin bütünüdür. Bu ayrım sayesinde mobil cihaz kullanıcı dostu ve hafif kalırken, ağır model çıkarımı kontrollü bir sunucu ortamında yapılabilir. Ayrıca aynı kaynak video gelecekte daha iyi model sürümleriyle yeniden işlenebilir.",
            "Projenin hedef kullanıcıları üç grupta düşünülmüştür. İlk grup oyun ve animasyon prototipi geliştiren öğrenci ve bağımsız geliştiricilerdir. Bu kullanıcılar pahalı mocap stüdyosuna erişmeden hızlı hareket verisi üretmek ister. İkinci grup küçük içerik üretim ekipleridir; onlar için kullanım kolaylığı, export formatı ve hızlı sonuç önemlidir. Üçüncü grup ise araştırma ve eğitim ortamlarıdır; burada sistemin ürettiği raporlar, kalite metrikleri ve tekrar işlenebilir video arşivi önem kazanır.",
            "Mocap'ın mevcut hali bir ticari ürünün nihai sürümü olarak değil, üretime taşınabilir bir mühendislik prototipi olarak konumlandırılmalıdır. Bu ayrım raporda açık biçimde yapılmalıdır. Proje çalışır mimari kararları, veri modellerini, upload ve processing süreçlerini, kalite raporu üretimini ve export dosyası akışını göstermektedir. Buna karşılık üretim seviyesinde kullanıcı hesabı, ödeme sistemi, KVKK/GDPR uyumlu veri saklama politikası, ölçekli GPU maliyet yönetimi ve gerçek çoklu kamera solve gibi konular sonraki faz geliştirmeleri olarak ele alınmalıdır.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 1.2. Hedef kullanıcı ve değer önerisi matrisi", caption_style)
    add_table(
        document,
        ["Kullanıcı grubu", "Temel ihtiyacı", "Mocap değer önerisi"],
        [
            ["Öğrenci/geliştirici", "Kısa sürede animasyon prototipi üretmek.", "Telefon kamerası ile çekim, backend solve, BVH export ve kalite raporu."],
            ["Bağımsız oyun ekibi", "Düşük maliyetli karakter hareketi oluşturmak.", "Stüdyo kurmadan tekrar işlenebilir video ve DCC uyumlu çıktı."],
            ["Küçük animasyon stüdyosu", "Hızlı referans hareketi ve ön üretim verisi almak.", "Proje/take yapısı, export listesi, kalite uyarıları ve arşivlenebilir artifact seti."],
            ["Eğitim/araştırma", "Mocap pipeline bileşenlerini gözlemlemek.", "Metadata, job timeline, pipeline report ve test edilebilir modüler mimari."],
        ],
        [1.8, 2.1, 2.8],
    )
    add_caption(document, "Tablo 1.3. Proje kapsamı ve kapsam dışı kararlar", caption_style)
    add_table(
        document,
        ["Alan", "Kapsam içinde", "Bu aşamada kapsam dışında"],
        [
            ["Mobil kayıt", "Native kamera preview, video kaydı, capture metadata, yerel take kaydı.", "Cihaz üzerinde tam WHAM/SMPL inference."],
            ["Backend", "Proje, take, upload, processing job ve export yönetimi.", "Tam üretim kimlik doğrulama ve ödeme sistemi."],
            ["Worker", "Video normalizasyonu, WHAM/SMPL solve, cleanup, BVH ve rapor artifact'leri.", "Tüm kamera açılarından nihai true multi-view optimizasyon."],
            ["Güvenlik", "Signed URL, object size kontrolü, credential maskeleme, veri bütünlüğü mantığı.", "Tam KVKK/GDPR yaşam döngüsü otomasyonu ve kurumsal DLP."],
            ["Ticarileşme", "MVP konumu, hedef pazar, maliyet kalemleri, ürünleştirme yol haritası.", "Canlı ödeme altyapısı ve resmi şirketleşme süreci."],
        ],
        [1.4, 2.6, 2.7],
    )


def add_system_expansion(document: Document, body_style: str, caption_style: str) -> None:
    add_paragraphs(
        document,
        [
            "Sistemin mobil tarafında veri modeli tasarımının merkezinde take kavramı bulunur. Take, kullanıcının bir hareket denemesini temsil eder ve yalnızca video dosyasından ibaret değildir. Take içinde proje ilişkisi, capture mode, yerel dosya yolu, çekim süresi, remote upload state, backend take id, processing job id ve export sonucu gibi alanlar bir araya gelir. Bu yapı sayesinde kullanıcı uygulamadan çıkıp geri geldiğinde çekim geçmişi korunabilir ve upload/processing durumu tekrar okunabilir.",
            "Capture metadata, projenin hareket yakalama hattı için en kritik sözleşmelerinden biridir. Metadata içinde çekim modu, device id, device index, camera role, timestamp, video süresi, çözünürlük, fps, orientation, quality flags ve çoklu kamera senkronizasyon bilgileri bulunur. Worker tarafında model çalıştırmadan önce bu metadata'nın eksiksiz ve tutarlı olması beklenir. Eksik metadata, yanlış model kullanımından çok daha erken fark edilmelidir; çünkü yanlış veya eksik metadata ile üretilen hareket çıktısı kullanıcıya güvenilir görünebilir ama teknik olarak yanıltıcı olabilir.",
            "Upload mimarisi özellikle büyük video dosyaları için tasarlanmıştır. Mobil uygulama önce backend'den upload session ve imzalı URL'ler alır. Metadata ve video ayrı object key'lere yüklenir. Upload tamamlandıktan sonra backend'e complete çağrısı yapılır. Backend object storage üzerinde dosya varlığını ve boyutu doğrular. Bu yöntem API sunucusunu dosya aktarım trafiğinden korur, hata durumlarında yeniden denemeyi kolaylaştırır ve worker'ın kaynak videoyu deterministik object key üzerinden bulmasını sağlar.",
            "Processing job akışı, backend ile worker arasındaki üretim sözleşmesidir. Job oluşturulduğunda take'in işlenebilir durumda olması gerekir. Worker işi aldığında önce kaynakları indirir, sonra video normalizasyonu yapar, daha sonra model çıkarımını çalıştırır. İşlem sırasında progress, currentStage, errorCode ve metrics alanları güncellenir. Kullanıcı tarafındaki ProcessingStatus ekranı bu state bilgisini okuyarak bekleme, hata, iptal ve sonuç durumlarını gösterir.",
            "Worker pipeline içinde video normalizasyonu hafife alınmaması gereken bir aşamadır. Mobil cihazlardan gelen videolar farklı codec, rotation metadata, variable frame rate veya container özellikleri taşıyabilir. WHAM ve benzeri modeller için tutarlı FPS, çözünürlük, yön ve frame erişimi önemlidir. Bu nedenle FFmpeg tabanlı normalize aşaması, model kalitesinden önce veri hazırlama kalitesini güvence altına alır. Rapor içinde normalize edilmiş video ile kaynak video arasındaki süre, fps ve dosya boyutu farkı gösterilmelidir.",
            "Export sistemi kullanıcıya tek bir dosya değil, bir artifact seti sunar. BVH dosyası DCC araçlarına aktarım için gereklidir. solved_motion JSON modelden çıkan hareket temsilini saklar. SMPL parameters dosyası vücut modeli parametrelerini taşır. quality_report teknik kaliteyi ve uyarıları açıklar. motion_pipeline_report hangi aşamaların çalıştığını, hangi artifact'lerin üretildiğini ve final animasyon kaynağının ne olduğunu gösterir. Bu artifact seti, projenin test edilebilir ve açıklanabilir olmasını sağlar.",
            "Çoklu kamera tarafında mimari, final animasyonu otomatik olarak iyileştiren sihirli bir katman gibi sunulmamalıdır. Projede multi-view session yönetimi, cihaz rolleri, join token, frame sync, kalibrasyon ve triangulation için altyapı vardır. Bu altyapı diagnostic raporlar ve ileride gerçek çoklu kamera kısıtlarının model optimizasyonuna bağlanması için önemlidir. Mevcut raporda bu durum açıkça belirtilmeli; başarı iddiası gerçek çalışan tek kamera WHAM hattı ve çoklu kamera tanısal verileri üzerinden kurulmalıdır.",
            "Yazılım mimarisinde hata yönetimi özel önem taşır. API yanıtlarında yalnızca genel hata mesajı dönmek yerine error code, request id ve retryable bilgisi verilmesi kullanıcı arayüzünü daha güvenilir yapar. Örneğin upload sırasında signed URL süresi dolarsa mobil uygulama fresh URL alarak tekrar deneyebilir. Buna karşılık metadata şeması hatalıysa retry anlamlı değildir; kullanıcıya çekimi tekrar yapması veya uygulama sürümünü güncellemesi gerektiği anlatılmalıdır.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 2.5. Mobil ekranlar ve kullanıcı akışındaki görevleri", caption_style)
    add_table(
        document,
        ["Ekran/akış", "Görev", "Rapora eklenecek ayrıntı"],
        [
            ["Capture", "Kamera preview ve kayıt kontrolü.", "İzin durumu, mod seçimi, kayıt süresi ve native kamera entegrasyonu."],
            ["Motion Preview/Review", "Kayıt sonrası take inceleme.", "Yerel video önizleme, metadata özeti ve upload'a geçiş."],
            ["Upload Progress", "Video ve metadata yükleme sürecini gösterme.", "Yüzde ilerleme, retry, remote state ve hata ekranları."],
            ["Processing Status", "Backend job state takibi.", "queued, running, failed, succeeded durumları ve polling davranışı."],
            ["Export Result", "Üretilen artifact dosyalarını gösterme.", "BVH, quality report, preview ve indirme bağlantıları."],
            ["Projects/Takes", "Kullanıcının kayıt geçmişini organize etme.", "Proje altında take listesi, yerel/remote durum ayrımı."],
        ],
        [1.5, 2.3, 3.0],
    )
    add_caption(document, "Tablo 2.6. API uçları ve sistem sözleşmeleri", caption_style)
    add_table(
        document,
        ["API grubu", "Temel sorumluluk", "Başarı kanıtı"],
        [
            ["Projects", "Kullanıcının çalışma alanlarını oluşturmak ve listelemek.", "Project id, name ve createdAt alanlarıyla kayıt."],
            ["Takes", "Çekim kayıtlarını backend tarafında temsil etmek.", "captureMode, expectedVideoCount ve status alanları."],
            ["Capture sessions", "Dual/pro cihaz eşleşmesini yönetmek.", "joinToken, deviceRoles, session status ve sync metadata."],
            ["Uploads", "Signed URL üretmek ve object doğrulamak.", "metadata/video object key, file size ve complete state."],
            ["Process", "Worker job başlatmak ve durum okumak.", "processingJobId, progress, currentStage, terminal state."],
            ["Exports", "Üretilen dosyaları listelemek ve indirme URL sağlamak.", "format, artifactName, storageKey ve signed download URL."],
        ],
        [1.4, 2.6, 2.8],
    )
    add_caption(document, "Tablo 2.7. Worker işlem hattı aşamaları", caption_style)
    add_table(
        document,
        ["Aşama", "Girdi", "Çıktı/rapor"],
        [
            ["Ingest", "capture_videos kayıtları ve S3 object key'leri.", "Kaynak dosya çalışma dizinine indirilir."],
            ["Normalize", "Kaynak video dosyası.", "FFmpeg normalize edilmiş video ve video_info."],
            ["Multi-view diagnostic", "Birden fazla video, sync ve calibration metadata.", "multi_view_quality_report ve reconstruction uyarıları."],
            ["WHAM/SMPL solve", "Normalize video ve model asset'leri.", "solved_motion, smpl_parameters ve overlay preview."],
            ["Cleanup", "Ham hareket verisi.", "cleanup_report, filtrelenmiş motion ve root stabilization bilgisi."],
            ["BVH export", "Temizlenmiş iskelet/motion verisi.", "BVH dosyası, export validation ve import smoke sonucu."],
            ["Report publish", "Tüm stage metrikleri ve artifact referansları.", "quality_report ve motion_pipeline_report."],
        ],
        [1.5, 2.5, 2.8],
    )
    add_caption(document, "Tablo 2.8. Artifact sözleşmesi ve kullanım amacı", caption_style)
    add_table(
        document,
        ["Artifact", "Format", "Kullanım amacı"],
        [
            ["source_video", "mp4/mov", "Tekrar işleme ve kalite karşılaştırması için orijinal kaynak."],
            ["normalized_video", "mp4", "Model runtime için standartlaştırılmış giriş."],
            ["solved_motion", "json", "Eklem pozisyonları, root motion ve frame bazlı hareket verisi."],
            ["smpl_parameters", "json/npz", "SMPL pose, shape ve camera/model parametreleri."],
            ["quality_report", "json", "Teknik kalite skoru, uyarılar, hatalar ve validasyon sonucu."],
            ["motion_pipeline_report", "json", "Stage süreleri, final animation source ve artifact lineage."],
            ["animation_bvh", "bvh", "Blender, Maya, MotionBuilder, Unity/Unreal pipeline için aktarım dosyası."],
            ["preview_overlay", "mp4", "Kullanıcıya hareket çözümünün görsel kontrolünü sunan önizleme."],
        ],
        [1.7, 1.0, 4.1],
    )
    add_paragraphs(
        document,
        [
            "Veri yaşam döngüsü raporda ayrıca gösterilmelidir; çünkü Mocap'ta aynı hareket denemesi farklı aşamalarda farklı veri türlerine dönüşür. İlk aşamada kullanıcı cihazında ham video ve local take metadata'sı vardır. İkinci aşamada backend, take ve upload session kayıtlarıyla bu veriyi remote sistemde temsil eder. Üçüncü aşamada worker, videoyu normalize eder ve model runtime için ara dosyalar üretir. Dördüncü aşamada export artifact'leri ve kalite raporları oluşur. Beşinci aşamada kullanıcı bu dosyaları indirir veya saklama süresi sonunda sistem veriyi temizler.",
            "Bu yaşam döngüsünün yazılması iki nedenle önemlidir. Birincisi, veri güvenliği değerlendirmesi hangi dosyanın nerede saklandığını bilmeden yapılamaz. İkincisi, ticari ürün planında depolama ve GPU maliyeti ancak hangi dosyanın ne kadar süre tutulacağı bilindiğinde hesaplanabilir. Örneğin orijinal video uzun süre saklanırsa tekrar işleme avantajı artar; ancak depolama maliyeti ve gizlilik sorumluluğu da artar. Kısa saklama süresi maliyeti düşürür fakat kullanıcı geçmiş çekimleri yeniden işleyemez.",
            "Veritabanı migration'ları da sistematik yaklaşım içinde değerlendirilmelidir. projects, takes, capture_sessions, capture_videos, processing_jobs ve export_files tabloları arasındaki ilişkiler zamanla genişleyebilir. Yeni bir export türü, yeni bir capture mode veya yeni bir quality metric eklendiğinde migration dosyasının hem mevcut veriyi bozmaması hem de worker ve mobil uygulamanın beklediği alanları tutarlı şekilde sağlaması gerekir.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 2.9. Veri yaşam döngüsü ve saklama kararları", caption_style)
    add_table(
        document,
        ["Aşama", "Veri", "Saklama/işleme kararı"],
        [
            ["Cihaz üstü kayıt", "Ham video, local take, capture metadata.", "Upload tamamlanana kadar cihazda tutulur; kullanıcı isterse yerel kopya kalabilir."],
            ["Upload", "metadata.json ve kaynak video object'leri.", "Signed URL ile yüklenir; object size ve metadata doğrulanır."],
            ["İşleme", "Normalize video, frame cache, model ara çıktıları.", "Worker çalışma alanında geçici tutulur; job sonunda temizlenmelidir."],
            ["Export", "BVH, solved_motion, quality_report, pipeline_report.", "Kullanıcı indirmesi ve tekrar görüntüleme için belirlenen retention süresince saklanır."],
            ["Log/metric", "Job state, stage süreleri, hata kodları.", "Destek ve QA için tutulur; hassas URL/secret içermemelidir."],
            ["Silme", "Video ve artifact nesneleri.", "Kullanıcı talebi veya retention süresi sonunda object storage ve DB kayıtları uyumlu temizlenmelidir."],
        ],
        [1.5, 2.5, 2.8],
    )


def add_experiment_expansion(document: Document, body_style: str, caption_style: str) -> None:
    add_paragraphs(
        document,
        [
            "Deney düzeneğinde ölçülmesi gereken ilk veri çekim kalitesidir. Çekim kalitesi yalnızca videonun açılıp açılmadığıyla ölçülmemelidir. Kadrajda tüm vücudun görünmesi, oyuncunun kameraya göre konumu, zemin temasının takip edilebilirliği, arka plan karmaşıklığı, ışık seviyesi, hareket hızı ve kamera stabilitesi değerlendirilmelidir. Bu veriler quality_report ile birlikte rapora işlendiğinde başarısız sonuçların nedenleri daha iyi açıklanabilir.",
            "İkinci ölçüm grubu upload ve backend güvenilirliğidir. Aynı video iyi bir kamerayla çekilmiş olsa bile upload sırasında dosya eksik giderse veya metadata yanlış eşleşirse worker doğru sonuç üretemez. Bu nedenle deneylerde metadata upload, video upload, upload complete, processing create ve export list adımlarının her biri ayrı kanıtla gösterilmelidir. Backend kayıtlarından object key, file size, state ve job id alanları maskelenmiş biçimde rapora eklenebilir.",
            "Üçüncü ölçüm grubu model ve export kalitesidir. WHAM/SMPL solve sonucunda yalnızca dosya oluşması yeterli değildir; dosyanın içeriği de doğrulanmalıdır. Solved motion içinde kare sayısı beklenen video süresiyle tutarlı olmalı, eklem rotasyonlarında NaN veya Infinity bulunmamalı, root motion aşırı sıçrama göstermemeli ve BVH dosyası DCC aracı tarafından import edilebilmelidir. Raporda Blender import ekran görüntüsü bu nedenle güçlü bir kanıttır.",
            "Çoklu kamera deneylerinde ana kabul kriteri final animasyon kalitesinin mutlaka artması değil, sistemin çoklu cihaz verisini doğru toplayıp tanısal bilgi üretebilmesidir. Session'a beklenen cihaz sayısı katılıyor mu, device role değerleri doğru mu, frame eşleşmesi yapılıyor mu, sync confidence hesaplanıyor mu, calibration eksikse sistem bunu uyarı olarak belirtiyor mu gibi sorular cevaplanmalıdır. Böylece çoklu kamera altyapısının gerçek durumu dürüst ve ölçülebilir şekilde sunulur.",
            "Sanal laboratuvar testlerinde fixture veri setleri kullanılmalıdır. Örneğin kısa bir solo video, bozuk metadata örneği, eksik video upload senaryosu, yanlış capture mode değeri, missing calibration multi-view paketi ve küçük bir synthetic solved motion çıktısı hazırlanabilir. Bu fixture'lar sayesinde her değişiklikten sonra backend ve worker davranışları tekrar edilebilir biçimde kontrol edilir. Raporun eklerinde bu fixture senaryolarının isimleri ve beklenen sonuçları verilmelidir.",
            "Test sonuçlarının raporda inandırıcı olması için başarı ve başarısızlık örnekleri birlikte gösterilmelidir. Başarılı bir solo çekimin yanında, düşük ışık veya eksik tüm vücut görünürlüğü nedeniyle kalite uyarısı üreten bir örnek de eklenmelidir. Bu yaklaşım projenin sadece güzel çalışan bir demo videosu üzerinden değil, hata durumlarını da tanıyan bir mühendislik sistemi üzerinden değerlendirildiğini gösterir.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 3.4. Önerilen deney veri seti", caption_style)
    add_table(
        document,
        ["Veri seti", "Amaç", "Beklenen rapor çıktısı"],
        [
            ["Solo iyi koşul", "Ana üretim hattının başarılı çalışmasını göstermek.", "Job succeeded, BVH export, quality grade yüksek."],
            ["Solo düşük ışık", "Model kalite uyarılarını gözlemlemek.", "quality_report warnings ve düşük confidence açıklaması."],
            ["Solo hızlı hareket", "Motion smoothing ve root stability etkisini görmek.", "cleanup_report ve jitter/velocity uyarıları."],
            ["Eksik kadraj", "Kısmi görünmezlik durumunu test etmek.", "tracking warning, quality score düşüşü."],
            ["Dual kamera", "Session, sync ve diagnostic akışı doğrulamak.", "matched frame count, sync confidence, diagnostic report."],
            ["Pro dört kamera", "Cihaz rolü ve beklenen video sayısı mantığını test etmek.", "expectedDeviceCount=4 ve upload readiness davranışı."],
            ["Bozuk metadata", "Validator ve hata yönetimini test etmek.", "Job başlatılmadan açıklayıcı hata."],
        ],
        [1.5, 2.5, 2.8],
    )
    add_caption(document, "Tablo 3.5. QA kanıt matrisi", caption_style)
    add_table(
        document,
        ["Kanıt", "Nasıl alınır", "Neyi kanıtlar"],
        [
            ["Ekran görüntüsü", "Mobil uygulamadan capture/upload/processing/export ekranları.", "Kullanıcı akışının uçtan uca varlığını."],
            ["Backend kayıt örneği", "PostgreSQL sorgusu veya API yanıtı.", "Veri modelinin doğru dolduğunu."],
            ["Object storage listesi", "MinIO/S3 bucket içeriği.", "Video ve artifact dosyalarının depolandığını."],
            ["Worker log özeti", "Stage süreleri ve terminal state.", "Model pipeline'ın hangi aşamalardan geçtiğini."],
            ["Quality report", "Üretilen JSON rapor.", "Teknik kalite ve uyarıların hesaplandığını."],
            ["BVH import", "Blender import ekranı veya smoke test.", "Export dosyasının DCC akışında kullanılabilir olduğunu."],
        ],
        [1.5, 2.5, 2.8],
    )


def add_security_expansion(document: Document, body_style: str, caption_style: str) -> None:
    add_paragraphs(
        document,
        [
            "Video verisinin kişisel veri niteliği taşıyabileceği kabul edilmelidir. Bir kullanıcının yüzü, vücut yapısı, hareket biçimi, çekim ortamı ve zaman bilgisi video içinde bulunabilir. Bu nedenle raporda örnek veri paylaşılırken gerçek kullanıcı görüntüleri yerine izinli veya anonimleştirilmiş test görüntüleri kullanılmalıdır. Eğer gerçek kullanıcı videosu kullanılacaksa açık rıza, saklama süresi, silme politikası ve erişim yetkileri ayrıca tanımlanmalıdır.",
            "Object storage mimarisinde dosyaların erişim politikası kritik önemdedir. Bucket'ların herkese açık olmaması gerekir. Mobil uygulama video yüklemek için yalnızca belirli object key'e, belirli süre geçerli imzalı URL ile erişmelidir. Download tarafında da export dosyaları için geçici URL üretilmeli, bu URL'ler log dosyalarına veya rapor ekran görüntülerine yazılmamalıdır. Rapor kanıtlarında URL'ler maskelenmelidir.",
            "Backend tarafında üretim öncesi en önemli güvenlik açığı kimlik doğrulama ve yetkilendirme katmanıdır. Geliştirme ortamında bearer token'ın kullanıcı id gibi yorumlanması hızlı prototip için yeterli olabilir; ancak ticari veya çok kullanıcılı ortamda bu yaklaşım kabul edilmemelidir. Üretim sürümünde JWT/OAuth tabanlı doğrulama, token imzası, token süresi, refresh mekanizması, kullanıcı-take sahipliği kontrolü ve rol tabanlı erişim eklenmelidir.",
            "Veri bütünlüğü açısından hash veya checksum kullanımı sonraki faz için önerilmelidir. Mevcut object size doğrulaması eksik upload'ları yakalayabilir; ancak dosyanın içerik bütünlüğünü tam kanıtlamaz. İleri aşamada metadata ve video için SHA-256 checksum alanı eklenebilir. Mobil uygulama upload öncesi checksum üretir, backend complete aşamasında object metadata veya worker indirimi sonrasında checksum doğrulaması yapar. Bu sayede bozuk dosya processing hattına girmeden engellenebilir.",
            "Gizlilik açısından raporun kendisi de güvenlik kontrolünden geçirilmelidir. Ekran görüntülerinde bucket adı, object key'in kullanıcıya özel bölümü, signed URL query parametreleri, access key, secret key, veritabanı URL'si, RunPod API key, model path, kullanıcı id ve cihaz id gibi alanlar gizlenmelidir. Bu bilgiler doğrudan uygulamanın güvenlik sınırlarını açığa çıkarabilir.",
            "Çoklu kamera akışında veri güvenliği yalnızca dosya erişimiyle sınırlı değildir. Session join token yetkisiz bir kişi tarafından ele geçirilirse aynı capture session'a yanlış cihaz katılabilir. Bu nedenle join token'ın kısa süreli olması, session status geçişlerinin kontrol edilmesi, expected device count dışındaki katılımların reddedilmesi ve cihaz rolü çakışmalarının backend tarafında engellenmesi gerekir.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 5.3. Kişisel veri ve gizlilik kontrol listesi", caption_style)
    add_table(
        document,
        ["Veri türü", "Risk", "Kontrol/öneri"],
        [
            ["Kaynak video", "Kullanıcı yüzü, beden hareketi ve çekim ortamı görünebilir.", "Açık rıza, özel bucket, geçici URL, saklama süresi ve silme isteği."],
            ["Capture metadata", "Cihaz, zaman ve çekim bağlamı çıkarılabilir.", "Gereksiz alanları toplamama, maskeleme, erişim kontrolü."],
            ["Model çıktısı", "Kullanıcının hareket profili tekrar oluşturulabilir.", "Kullanıcı sahipliği, signed download, retention policy."],
            ["Log kayıtları", "URL, object key veya hata içinde hassas değer sızabilir.", "Structured logging, secret redaction, request id kullanımı."],
            ["Rapor kanıtları", "Ekran görüntüsünde token veya dosya yolu görünebilir.", "Teslim öncesi görsel maskeleme ve metadata temizliği."],
        ],
        [1.5, 2.6, 2.7],
    )
    add_caption(document, "Tablo 5.4. Üretim güvenliği için önerilen ek kontroller", caption_style)
    add_table(
        document,
        ["Kontrol", "Açıklama", "Öncelik"],
        [
            ["Gerçek auth", "JWT/OAuth, token imzası, refresh ve kullanıcı sahipliği kontrolü.", "Yüksek"],
            ["Rate limit", "Upload init, process create ve signed download endpoint'lerinde kötüye kullanımı azaltır.", "Yüksek"],
            ["Checksum", "Video ve metadata bütünlüğünü object size dışında doğrular.", "Orta"],
            ["Audit log", "Kim hangi take/export dosyasına erişti izlenir.", "Orta"],
            ["Retention policy", "Video ve artifact dosyalarının ne kadar saklanacağı belirlenir.", "Yüksek"],
            ["DLP/redaction", "Log ve raporlarda secret/signed URL görünmesini engeller.", "Yüksek"],
            ["CORS hardening", "Üretim alan adı dışındaki istekleri sınırlar.", "Orta"],
        ],
        [1.5, 3.0, 1.3],
    )


def add_change_management_section(
    document: Document,
    h1_style: str,
    h2_style: str,
    first_style: str,
    body_style: str,
    caption_style: str,
) -> None:
    add_heading(document, "DEĞİŞİKLİK YÖNETİMİ VE MÜHENDİSLİK STANDARTLARI", h1_style)
    add_heading(document, "Mocap Yazılım Geliştirme Süreci", h2_style)
    add_paragraph(
        document,
        "Mocap projesinde değişiklik yönetimi, mobil uygulama, backend API, veritabanı migration'ları, worker pipeline ve model runtime gibi birbirine bağlı bileşenlerde yapılacak güncellemelerin kontrollü ilerletilmesini ifade eder. Hareket yakalama sistemi çok bileşenli olduğu için küçük görünen bir değişiklik bile tüm zinciri etkileyebilir. Örneğin capture metadata alanının mobilde değiştirilmesi backend validator'ını, worker'ın input okuma mantığını, quality_report şemasını ve mobil sonuç ekranını aynı anda etkileyebilir. Bu nedenle proje geliştirilirken her değişiklik izlenebilir, test edilebilir ve geri alınabilir şekilde ele alınmalıdır.",
        first_style,
        first=True,
    )
    add_paragraphs(
        document,
        [
            "Değişiklik yönetiminin ilk adımı değişiklik talebinin açık tanımlanmasıdır. Talep; problem, gerekçe, etkilenen bileşenler, beklenen davranış, kabul kriteri ve risk alanlarıyla birlikte yazılmalıdır. Örneğin 'dual kamera sync iyileştirilecek' ifadesi yeterli değildir. Bunun yerine hangi sync metriğinin iyileştirileceği, hangi capture mode'larda geçerli olacağı, worker output şemasına yeni alan eklenip eklenmeyeceği ve mobil arayüzde hangi değerin gösterileceği belirtilmelidir.",
            "İkinci adım etki analizidir. Mocap mimarisinde etki analizi en az beş katmanda yapılmalıdır: mobil kullanıcı akışı, domain model, backend API sözleşmesi, veritabanı şeması ve worker artifact sözleşmesi. Bu analiz yapılmadan doğrudan kod değişikliği yapmak, özellikle metadata ve export şemalarında geriye dönük uyumluluk sorunlarına neden olabilir. Rapor içinde bir değişikliğin hangi katmanları etkilediğini gösteren matris sunulmalıdır.",
            "Üçüncü adım sürümleme ve geriye dönük uyumluluktur. Mobil uygulama ve backend aynı anda güncellenmeyebilir. Bu nedenle API contract ve metadata schema alanlarında version bilgisi kullanılmalıdır. CaptureMetadata içinde schemaVersion benzeri alanlar korunmalı, worker eski ve yeni sürüm arasında anlamlı hata üretebilmelidir. Semantik sürümleme yaklaşımıyla major değişiklikler kırıcı sözleşme değişikliği, minor değişiklikler geriye uyumlu yeni özellik, patch değişiklikler hata düzeltmesi olarak sınıflandırılabilir.",
            "Dördüncü adım test kapısıdır. Değişiklik main geliştirme hattına alınmadan önce ilgili birim testleri, backend endpoint testleri, metadata validation testleri, worker smoke testleri ve seçilmiş gerçek cihaz senaryoları çalıştırılmalıdır. Video tabanlı sistemlerde yalnızca statik tip kontrolü yeterli değildir; en az bir kısa solo video ve bir hata senaryosu ile uçtan uca test yapılması gerekir. Değişiklik sonucunda üretilen quality_report ve motion_pipeline_report eski fixture sonuçlarıyla karşılaştırılmalıdır.",
            "Beşinci adım release kararıdır. Release öncesinde değişiklik notu, migration planı, rollback planı, bilinen kısıtlar ve kullanıcıya etkisi yazılmalıdır. Örneğin worker artifact şeması değiştiyse eski export result ekranının bu dosyaları okuyup okuyamayacağı kontrol edilmelidir. Veritabanı migration'ı eklenmişse migration'ın tekrar çalıştırıldığında bozulmaması ve rollback ihtiyacı durumunda veri kaybı yaratmaması değerlendirilmelidir.",
            "Mühendislik standartları açısından proje gereksinim yönetimi, yazılım yaşam döngüsü, güvenli API geliştirme, sürümleme, dokümantasyon ve test tekrarlanabilirliği ilkelerine dayandırılmalıdır. Gereksinimler açık, ölçülebilir ve test edilebilir olmalıdır. Örneğin 'sistem hızlı olmalıdır' yerine '10 saniyelik 1080p solo video için upload, işleme ve export süreleri ayrı metrikler olarak raporlanmalıdır' ifadesi daha mühendislik odaklıdır.",
            "Kod kalitesi standardı olarak tip güvenliği, modülerlik, servis/repository ayrımı, domain model sözleşmeleri ve açık hata tipleri korunmalıdır. Mobil tarafta feature tabanlı klasörleme, backend tarafında route-service-repository ayrımı, worker tarafında stage bazlı pipeline yapısı projenin bakımını kolaylaştırır. Bu mimari düzen raporda yalnızca dosya isimleriyle değil, değişiklik yönetimine nasıl katkı sağladığıyla birlikte açıklanmalıdır.",
            "Veri standardı açısından JSON şemaları, migration dosyaları ve artifact adları kararlı tutulmalıdır. Worker tarafından üretilen her artifact için format, storage key, dosya boyutu, üretim stage'i ve kullanıcı arayüzünde gösterim biçimi tanımlanmalıdır. Bu yaklaşım hem QA sürecini hem de ticarileşme aşamasında müşteri destek süreçlerini kolaylaştırır.",
            "Güvenlik standardı açısından OWASP API güvenliği ve uygulama güvenliği ilkeleri referans alınmalıdır. Özellikle authorization, broken object level authorization, rate limiting, excessive data exposure ve logging konuları Mocap için önemlidir. Çünkü kullanıcıya ait videolar ve export dosyaları object key üzerinden yönetilmektedir. Üretim seviyesinde her take ve export dosyası için sahiplik kontrolü yapılmalıdır.",
            "Dokümantasyon standardı açısından her büyük değişiklikten sonra üç belge güncellenmelidir: kullanıcıya dönük davranış açıklaması, geliştiriciye dönük API/veri modeli dokümanı ve QA/test kanıtı. Bu üçlü güncellenmezse proje ilerledikçe çalışan kod ile rapor veya teknik dokümanlar arasında tutarsızlık oluşur. Bitirme kitapçığında bu yaklaşım, projenin sürdürülebilir mühendislik bakışıyla geliştirildiğini gösterecektir.",
        ],
        body_style,
    )
    add_heading(document, "Sürümleme, Etki Analizi ve Değişiklik Kontrolü", h2_style)
    add_diagram_box(
        document,
        [
            "Degisiklik talebi",
            "   v",
            "Etki analizi: Mobil | API | DB | Worker | Artifact",
            "   v",
            "Uygulama + test + fixture karsilastirma",
            "   v",
            "Review + release notu + rollback plani",
        ],
    )
    add_caption(document, "Şekil 8.1. Değişiklik yönetimi akışı", caption_style)
    add_caption(document, "Tablo 8.1. Değişiklik türleri ve kontrol adımları", caption_style)
    add_table(
        document,
        ["Değişiklik türü", "Örnek", "Zorunlu kontrol"],
        [
            ["Mobil UI değişikliği", "Processing ekranında yeni progress açıklaması.", "Ekran görüntüsü, metin taşması kontrolü, hata durumu testi."],
            ["Metadata şema değişikliği", "captureMode veya sync alanına yeni değer eklenmesi.", "Schema version, backend validator, worker input parser testi."],
            ["API contract değişikliği", "Upload complete yanıtına yeni alan eklenmesi.", "OpenAPI/endpoint dokümanı, geriye uyumluluk testi."],
            ["Veritabanı migration", "export_files tablosuna yeni kolon eklenmesi.", "Migration tekrar çalıştırma, rollback planı, mevcut kayıt kontrolü."],
            ["Worker pipeline değişikliği", "Cleanup filtresi veya BVH writer güncellemesi.", "Golden fixture, BVH import smoke test, quality report farkı."],
            ["Güvenlik değişikliği", "Signed URL TTL veya auth kontrolü değiştirme.", "Yetkisiz erişim testi, log redaction, kullanıcı sahipliği kontrolü."],
            ["Model runtime değişikliği", "WHAM checkpoint veya SMPL asset yolu güncellemesi.", "Runtime preflight, lisans kontrolü, eski fixture karşılaştırması."],
        ],
        [1.6, 2.4, 3.0],
    )
    add_heading(document, "Mühendislik Standartları ve Kalite Kapıları", h2_style)
    add_caption(document, "Tablo 8.2. Mühendislik standartları uyum matrisi", caption_style)
    add_table(
        document,
        ["Standart/ilke", "Mocap uygulaması", "Raporda gösterilecek kanıt"],
        [
            ["Gereksinim izlenebilirliği", "Her özellik kullanıcı ihtiyacı, sistem bileşeni ve test kanıtı ile bağlanır.", "Gereksinim-test matrisi."],
            ["Yazılım yaşam döngüsü", "Analiz, tasarım, geliştirme, test, release ve bakım aşamaları ayrılır.", "İş-zaman çizelgesi ve değişiklik kayıtları."],
            ["Semantik sürümleme", "API, metadata ve artifact şemalarında kırıcı değişiklik ayrımı yapılır.", "Sürüm notu ve schemaVersion alanı."],
            ["Güvenli API geliştirme", "Auth, authorization, rate limit, signed URL ve secret redaction planlanır.", "Güvenlik risk matrisi."],
            ["Test tekrarlanabilirliği", "Golden fixture, synthetic data ve gerçek cihaz smoke testleri korunur.", "QA kanıt matrisi ve test sonuçları."],
            ["Dokümantasyon", "API, veri modeli, worker artifact ve kullanıcı akışı güncel tutulur.", "Ek A çıktı listesi ve kaynakça."],
        ],
        [1.6, 3.0, 2.4],
    )
    add_heading(document, "Ekip Sorumlulukları ve Onay Mekanizması", h2_style)
    add_paragraphs(
        document,
        [
            "Değişiklik yönetiminde yalnızca teknik adımların değil, ekip sorumluluklarının da tanımlanması gerekir. Bitirme projesi birden fazla öğrencinin ortak çalışması olduğu için mobil uygulama, backend, worker, raporlama ve test sorumlulukları açık yazılmalıdır. Aksi durumda bir değişiklik tamamlandığında kimin test ettiği, kimin dokümante ettiği ve kimin final rapora işlediği belirsiz kalır.",
            "RACI benzeri basit bir sorumluluk matrisi bu belirsizliği azaltır. Responsible alanı işi yapan kişiyi, Accountable alanı son karardan sorumlu kişiyi, Consulted alanı teknik görüş veren kişiyi, Informed alanı ise değişiklikten haberdar edilmesi gereken kişiyi temsil eder. Akademik raporda isimler doğrudan yazılabilir veya ekip rolleriyle ifade edilebilir. Bu matrisi eklemek, projenin ekip çalışması ve mühendislik yönetimi açısından olgunlaştığını gösterir.",
            "Değişiklik kararları için önem derecesi de belirlenmelidir. Örneğin ekran metni değişikliği düşük risklidir; metadata şeması değişikliği orta-yüksek risklidir; worker model runtime veya veritabanı migration değişikliği yüksek risklidir. Yüksek riskli değişikliklerde mutlaka rollback planı, fixture testi ve rapor güncellemesi zorunlu tutulmalıdır.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 8.3. Değişiklik yönetimi sorumluluk matrisi", caption_style)
    add_table(
        document,
        ["Alan", "Responsible", "Accountable", "Consulted", "Informed"],
        [
            ["Mobil capture/review akışı", "Mobil geliştirici", "Proje ekibi", "Backend/QA", "Danışman ve ekip"],
            ["Backend API ve veri modeli", "Backend geliştirici", "Proje ekibi", "Mobil/worker", "Danışman ve ekip"],
            ["Worker ve motion pipeline", "Worker sorumlusu", "Proje ekibi", "Mobil/backend/QA", "Danışman ve ekip"],
            ["Test ve kalite kanıtları", "QA sorumlusu", "Proje ekibi", "Tüm teknik roller", "Danışman ve ekip"],
            ["Rapor/dokümantasyon", "Rapor sorumlusu", "Proje ekibi", "Tüm teknik roller", "Danışman ve jüri"],
            ["Ticarileşme ve güvenlik planı", "Ürün/analiz sorumlusu", "Proje ekibi", "Backend/güvenlik", "Danışman ve ekip"],
        ],
        [1.5, 1.4, 1.4, 1.4, 1.4],
    )
    add_heading(document, "Yayınlama, Doğrulama ve Geri Dönüş Planı", h2_style)
    add_diagram_box(
        document,
        [
            "Preflight -> Typecheck -> Backend QA -> Worker fixture -> Real device smoke",
            "                          v",
            "                    Release gate",
            "                          v",
            "             Version note + rollback plan",
        ],
    )
    add_caption(document, "Şekil 8.2. Release kapısı ve kalite güvence zinciri", caption_style)


def add_commercialization_section(
    document: Document,
    h1_style: str,
    h2_style: str,
    first_style: str,
    body_style: str,
    caption_style: str,
) -> None:
    add_heading(document, "TİCARİLEŞME PLANI", h1_style)
    add_heading(document, "MVP Kapsamı ve Değer Önerisi", h2_style)
    add_paragraph(
        document,
        "Mocap'ın ticarileşme potansiyeli, hareket yakalama teknolojisini pahalı stüdyo kurulumlarından çıkarıp mobil cihaz ve bulut işleme modeline yaklaştırmasından kaynaklanır. Ürünleştirme açısından proje, tek seferlik bir akademik demo yerine abonelik veya kullanım başına ödeme modeliyle sunulabilecek bir hareket yakalama servis altyapısı olarak değerlendirilebilir. Ticari plan hazırlanırken ürünün mevcut prototip seviyesi, teknik borçları, hedef pazar, maliyet kalemleri, gelir modeli, hukuki gereksinimler ve pazara çıkış adımları birlikte düşünülmelidir.",
        first_style,
        first=True,
    )
    add_paragraphs(
        document,
        [
            "Ürünün ilk MVP sürümü solo kamera hattı üzerine kurulmalıdır. Çünkü mevcut sistemde en güvenilir final animasyon yolu tek kamera WHAM/SMPL pipeline'ıdır. Kullanıcı bir video kaydeder, backend bu videoyu işler, kullanıcı BVH ve rapor çıktılarını alır. Dual/pro çoklu kamera özellikleri ilk ticari sürümde 'gelişmiş tanısal mod' veya 'beta özellik' olarak sunulabilir. Bu yaklaşım teknik riskleri azaltır ve kullanıcının gerçekten değer aldığı temel senaryoyu öne çıkarır.",
            "Değer önerisi üç ana noktaya dayanır. Birincisi erişilebilirliktir; özel kıyafet, marker veya stüdyo gerektirmez. İkincisi tekrar işlenebilirliktir; kaynak video saklandığı için model güncellemelerinde aynı take yeniden işlenebilir. Üçüncüsü açıklanabilirliktir; kullanıcı yalnızca dosya indirmez, kalite raporu ve pipeline raporu ile sonucun hangi sınırlara sahip olduğunu görür. Bu özellik ticari ürün güveni açısından önemlidir.",
        ],
        body_style,
    )
    add_heading(document, "Hedef Pazar ve Kullanım Senaryoları", h2_style)
    add_paragraphs(
        document,
        [
            "Hedef pazar öncelikle bağımsız oyun geliştiricileri, animasyon öğrencileri, küçük stüdyolar ve hızlı prototipleme ekipleridir. Bu gruplar için pahalı motion capture stüdyosu kiralamak veya marker tabanlı sistem almak ekonomik değildir. Mobil cihaz üzerinden çekim yapıp BVH çıktısı alabilmek, özellikle oyun prototipleri, kısa animasyonlar, eğitim materyalleri ve sosyal medya içerikleri için pratik değer sağlar.",
            "Pazara çıkış stratejisi teknik topluluklardan başlamalıdır. İlk kullanıcı grubu üniversite öğrencileri, game jam ekipleri, bağımsız Unity/Unreal geliştiricileri ve Blender kullanıcıları olabilir. Ürün tanıtımında 'telefonla çek, BVH olarak indir, Blender'da aç' gibi açık bir kullanım sonucu gösterilmelidir. Teknik blog yazısı, kısa demo videosu, GitHub dokümanları, Discord topluluğu ve eğitim videoları erken kullanıcı kazanımı için uygundur.",
            "Ürünün rekabet avantajı yalnızca video tabanlı mocap yapması değildir. Rakip veya alternatif araçlar karşısında Mocap'ın farkı, mobil uygulama ile backend pipeline'ı birlikte sunması, kalite raporu üretmesi, artifact lineage kaydı tutması ve çoklu kamera altyapısına genişleyebilir olmasıdır. Ancak bu avantajların ticari olarak anlamlı olması için final animasyon kalitesi, kullanıcı deneyimi ve export uyumluluğu düzenli olarak ölçülmelidir.",
        ],
        body_style,
    )
    add_diagram_box(
        document,
        [
            "MVP: Solo capture -> WHAM solve -> BVH export",
            "Growth: Pro presets -> batch jobs -> Unity/Unreal/Blender guides",
            "Scale: True multi-view -> team workspace -> API/enterprise controls",
        ],
    )
    add_caption(document, "Şekil 9.1. Ticarileşme yol haritası", caption_style)
    add_caption(document, "Tablo 9.1. Hedef pazar ve kullanım senaryoları", caption_style)
    add_table(
        document,
        ["Segment", "Kullanım senaryosu", "Satış mesajı"],
        [
            ["Oyun geliştirici", "Karakter prototipi için hızlı koşma, zıplama, dönüş hareketleri.", "Telefonla çek, BVH indir, oyun motorunda dene."],
            ["Animasyon öğrencisi", "Ders veya portfolyo için temel hareket referansı.", "Stüdyo gerektirmeden tekrar işlenebilir hareket verisi."],
            ["Küçük stüdyo", "Ön prodüksiyon ve referans animasyon.", "Düşük maliyetli hızlı çekim, kalite raporu ve export paketi."],
            ["Eğitim kurumu", "Mocap pipeline öğretimi ve öğrenci projeleri.", "Açıklanabilir raporlar, test senaryoları ve çoklu kamera altyapısı."],
            ["Araştırma/prototip ekibi", "Video tabanlı insan hareketi verisi toplama.", "API ve artifact seti üzerinden tekrarlanabilir deney."],
        ],
        [1.4, 2.8, 2.6],
    )
    add_heading(document, "Gelir Modeli ve Maliyet Yapısı", h2_style)
    add_paragraphs(
        document,
        [
            "Gelir modeli aşamalı tasarlanabilir. İlk model freemium olabilir: kullanıcı sınırlı süre, düşük çözünürlük veya sınırlı aylık işlem hakkıyla sistemi deneyebilir. İkinci model kullanım başına ödeme olabilir: her başarılı processing job veya her export paketi belirli kredi tüketir. Üçüncü model profesyonel abonelik olabilir: daha yüksek çözünürlük, batch processing, uzun video süresi, öncelikli GPU kuyruğu ve gelişmiş export formatları bu pakette sunulur.",
            "Maliyet tarafında en büyük değişken GPU işleme maliyetidir. Video dosyalarının depolanması ve indirilmesi de maliyet yaratır, ancak ağır WHAM/SMPL inference süresi ürün fiyatlandırmasını doğrudan etkiler. Bu nedenle ticarileşme planında ortalama video süresi, ortalama job süresi, GPU dakika maliyeti, depolama maliyeti, bandwidth maliyeti ve kullanıcı başına destek maliyeti hesaplanmalıdır. Kredi modeli bu maliyetleri kullanıcı davranışına göre dengelemek için uygundur.",
            "Operasyonel olarak destek ve hata yönetimi de ürünleştirme planına dahil edilmelidir. Kullanıcı başarısız bir job için ne görecek, kredi iadesi nasıl olacak, kalite raporu düşük çıktığında sistem bunu ücretlendirecek mi, upload yarıda kalırsa video ne kadar saklanacak, işleme kuyruğu yoğun olduğunda kullanıcıya ne gösterilecek gibi sorular ticarileşme açısından teknik kadar önemlidir.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 9.2. Gelir modeli ve maliyet kalemleri", caption_style)
    add_table(
        document,
        ["Model/kalem", "Açıklama", "Rapor için gerekli hesap"],
        [
            ["Freemium", "Sınırlı aylık job veya düşük çözünürlük.", "Ücretsiz kullanıcı başına ortalama GPU/storage maliyeti."],
            ["Kredi paketi", "Her başarılı processing job kredi harcar.", "Ortalama video süresi ve job başına maliyet."],
            ["Profesyonel abonelik", "Yüksek öncelik, uzun video, batch processing.", "Aylık aktif kullanıcı ve GPU kapasite planı."],
            ["Storage", "Kaynak video ve artifact saklama.", "GB/ay maliyeti ve retention süresi."],
            ["Bandwidth", "Video upload ve export download trafiği.", "Ortalama dosya boyutu ve indirme sayısı."],
            ["Destek", "Başarısız job, iade, kalite soruları.", "Destek talebi oranı ve operasyon zamanı."],
            ["Lisans/uyum", "Model asset, hukuki metin, gizlilik süreçleri.", "Tek seferlik kurulum ve sürekli bakım maliyeti."],
        ],
        [1.5, 2.6, 2.7],
    )
    add_heading(document, "Pazara Çıkış ve Ürünleştirme Yol Haritası", h2_style)
    add_paragraph(
        document,
        "MVP sonrası yol haritası üç faza ayrılabilir. İlk faz solo kamera üretim hattını sağlamlaştırır: auth, ödeme öncesi kullanım limiti, temel dashboard, Blender/Unity export doğrulaması ve gerçek cihaz QA. İkinci faz pro kullanıcı özelliklerini ekler: uzun video, batch jobs, gelişmiş cleanup presetleri, retargeting seçenekleri ve kalite karşılaştırma ekranı. Üçüncü faz ise true multi-view solve, ekip çalışma alanları, API erişimi ve kurumsal veri yönetimi özelliklerini hedefler.",
        body_style,
        first=True,
    )
    add_caption(document, "Tablo 9.3. Pazara çıkış ve ürünleştirme planı", caption_style)
    add_table(
        document,
        ["Faz", "Kapsam", "Başarı ölçütü"],
        [
            ["Faz 1 - Akademik MVP", "Solo capture, backend solve, BVH export, kalite raporu.", "En az 5 gerçek cihaz testinde uçtan uca başarılı job."],
            ["Faz 2 - Kapalı beta", "Auth, kullanıcı hesabı, kullanım limiti, hata iadesi mantığı.", "Beta kullanıcıların tekrar job oluşturma oranı."],
            ["Faz 3 - Creator paketi", "Blender/Unity presetleri, batch job, daha iyi export UX.", "Export indirme oranı ve düşük destek talebi."],
            ["Faz 4 - Pro kamera beta", "Dual/pro diagnostic iyileştirme ve calibration workflow.", "Sync confidence ve reconstruction metriklerinde ölçülebilir gelişme."],
            ["Faz 5 - Kurumsal/API", "Team workspace, API key, audit log, retention policy.", "Kurumsal pilot müşteri ve SLA hedefleri."],
        ],
        [1.5, 2.9, 2.4],
    )
    add_heading(document, "Risk Analizi ve Sürdürülebilirlik", h2_style)
    add_paragraphs(
        document,
        [
            "Ticarileşme öncesi hukuki ve etik gereksinimler tamamlanmalıdır. Kullanıcı videosu kişisel veri kabul edilerek açık rıza metni, aydınlatma metni, veri saklama süresi, silme talebi, kullanıcı hesabı kapatma ve üçüncü taraf GPU/depoma sağlayıcılarına veri aktarımı açıklanmalıdır. Türkiye'de KVKK; uluslararası kullanıcı hedeflenirse GDPR benzeri düzenlemeler dikkate alınmalıdır. Bu alan tamamlanmadan gerçek kullanıcı verisiyle ticari kullanım risklidir.",
            "Ticarileşme başarısı için ölçülmesi gereken metrikler kullanıcı sayısı, tamamlanan job sayısı, job başarı oranı, ortalama işleme süresi, kullanıcı başına GPU maliyeti, export indirme oranı, tekrar kullanım oranı, başarısız job nedenleri, destek talebi oranı ve kullanıcı başına gelir olarak belirlenebilir. Bu metrikler olmadan ürün fiyatlandırması ve teknik yatırım kararları tahmine dayalı kalır.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 9.4. Ticarileşme riskleri ve azaltma stratejileri", caption_style)
    add_table(
        document,
        ["Risk", "Etkisi", "Azaltma stratejisi"],
        [
            ["GPU maliyeti yüksekliği", "Kredi fiyatı ve karlılık baskılanır.", "Video süresi limiti, kuyruk önceliği, batch optimizasyonu."],
            ["Düşük animasyon kalitesi", "Kullanıcı güveni azalır.", "Quality report, çekim rehberi, başarısız job iade politikası."],
            ["Veri gizliliği riski", "Hukuki ve itibar kaybı.", "KVKK/GDPR metinleri, retention policy, delete request akışı."],
            ["Rakip araçlar", "Fiyat ve özellik baskısı.", "Mobil-first deneyim, raporlanabilir kalite, yerel pazar ve eğitim odaklı başlangıç."],
            ["Model lisansı", "Ticari kullanım sınırı oluşabilir.", "Lisans kontrolü, alternatif model araştırması, açık sözleşme yönetimi."],
            ["Çoklu kamera beklentisi", "Beta özelliğin final kalite gibi algılanması.", "Ürün metninde diagnostic/beta ayrımı ve ölçülebilir roadmap."],
        ],
        [1.5, 2.4, 3.0],
    )
    add_paragraphs(
        document,
        [
            "Ticarileşme planının daha anlaşılır olması için Lean Canvas benzeri bir özet de eklenmelidir. Bu özet, problemin, çözümün, hedef kitlenin, farklılaştırıcı değerin, gelir modelinin ve temel maliyetlerin tek tabloda görülmesini sağlar. Özellikle teknik projelerde ürün değeri bazen yalnızca mimari anlatım içinde kaybolur; canvas tablosu jüriye projenin pazara nasıl taşınacağını hızlıca gösterir.",
            "SWOT analizi ise ürünleştirme risklerini daha dengeli gösterir. Mocap'ın güçlü yönleri mobil erişilebilirlik ve uçtan uca pipeline iken, zayıf yönleri model runtime bağımlılığı, GPU maliyeti ve üretim seviyesi auth eksikliğidir. Fırsatlar oyun/animasyon içerik üretimi, eğitim pazarı ve bağımsız geliştirici topluluklarıdır. Tehditler ise rakip araçlar, veri gizliliği yükümlülükleri ve kullanıcıların kalite beklentisidir.",
            "12 aylık yol haritası ticarileşme planını somutlaştırır. Bu yol haritası akademik teslim sonrası ne yapılacağını gösterir: ilk üç ay teknik borçların kapatılması, sonraki üç ay kapalı beta, sonraki üç ay ücretli pilot ve son üç ay pro özellikler/kurumsal hazırlık olarak düzenlenebilir. Her dönem için ölçülebilir hedef konulmalıdır; aksi durumda yol haritası yalnızca niyet listesi olarak kalır.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 9.5. Lean Canvas özeti", caption_style)
    add_table(
        document,
        ["Alan", "Mocap için karşılığı", "Rapor notu"],
        [
            ["Problem", "Mocap stüdyoları pahalı, erişimi zor ve kurulum gerektirir.", "Giriş bölümündeki problem tanımıyla bağlanmalı."],
            ["Çözüm", "Mobil video kaydı, backend WHAM/SMPL solve, BVH ve rapor export.", "Bölüm 2 mimarisiyle desteklenmeli."],
            ["Hedef müşteri", "Bağımsız oyun geliştiricileri, animasyon öğrencileri, küçük stüdyolar.", "Tablo 7.1 ile ilişkilendirilmeli."],
            ["Değer önerisi", "Stüdyo kurmadan tekrar işlenebilir hareket yakalama.", "Kalite raporu ve artifact seti vurgulanmalı."],
            ["Kanallar", "Blender/Unity toplulukları, üniversiteler, game jam ekipleri, demo videoları.", "Pazara çıkış planına eklenmeli."],
            ["Gelir", "Freemium, kredi paketi, profesyonel abonelik.", "Tablo 9.2 maliyetleriyle dengelenmeli."],
            ["Maliyet", "GPU dakika, storage, bandwidth, destek ve hukuki uyum.", "Varsayımsal hesap Ek A'ya eklenmeli."],
            ["Metrikler", "Job başarı oranı, export indirme oranı, tekrar kullanım, GPU/job maliyeti.", "Pilot sonrası izlenecek KPI seti."],
        ],
        [1.4, 3.5, 2.0],
    )
    add_caption(document, "Tablo 9.6. SWOT analizi", caption_style)
    add_table(
        document,
        ["Kategori", "Mocap değerlendirmesi", "Stratejik sonuç"],
        [
            ["Güçlü yönler", "Mobil-first deneyim, backend pipeline, kalite raporu, BVH export.", "Demo ve MVP mesajı bu güçlü yönler üzerine kurulmalı."],
            ["Zayıf yönler", "GPU/runtime bağımlılığı, auth hardening ihtiyacı, true multi-view'in henüz beta olması.", "İlk ticari sürüm solo pipeline ile sınırlanmalı."],
            ["Fırsatlar", "Oyun geliştirme, eğitim, küçük stüdyolar, içerik üretimi, düşük maliyetli prototipleme.", "Topluluk odaklı erken kullanıcı kazanımı planlanmalı."],
            ["Tehditler", "Rakip mocap araçları, veri gizliliği regülasyonları, kalite beklentisi, model lisansı.", "Açık kalite raporu, hukuki uyum ve lisans kontrolü önceliklendirilmeli."],
        ],
        [1.4, 3.5, 2.0],
    )
    add_caption(document, "Tablo 9.7. 12 aylık ticarileşme yol haritası", caption_style)
    add_table(
        document,
        ["Dönem", "Odak", "Somut çıktı"],
        [
            ["0-3 ay", "Teknik borç ve üretim hazırlığı.", "Auth, rate limit, fixture set, gerçek cihaz QA, temiz deploy dokümanı."],
            ["3-6 ay", "Kapalı beta.", "20-50 kullanıcıyla solo capture pilotu, hata/iade akışı, kullanım limiti."],
            ["6-9 ay", "Ücretli creator paketi.", "Kredi sistemi denemesi, Blender/Unity rehberleri, export kalite dashboard'u."],
            ["9-12 ay", "Pro ve kurumsal hazırlık.", "Dual/pro beta, takım çalışma alanı, audit log, retention policy ve API erişimi."],
        ],
        [1.4, 3.4, 2.1],
    )


def add_appendix_expansion(document: Document, body_style: str, caption_style: str) -> None:
    add_paragraphs(
        document,
        [
            "Değişiklik yönetimi bölümü eklendiği için eklerde ayrıca bir değişiklik kayıt tablosu bulunmalıdır. Bu tabloda tarih, değişiklik adı, etkilenen bileşen, risk seviyesi, test kanıtı, onay durumu ve rollback notu alanları yer almalıdır. Rapor tesliminden önce son üç önemli proje değişikliği bu tabloya işlenirse jüri, projenin kontrollü mühendislik süreciyle geliştirildiğini görebilir.",
            "Ticarileşme planı için eklerde varsayımsal maliyet hesabı yer almalıdır. Bu hesap gerçek fatura olmak zorunda değildir; ancak bir job için ortalama video süresi, GPU işlem süresi, storage boyutu ve bandwidth varsayımı yapılmalıdır. Böylece freemium, kredi ve abonelik modellerinin teknik maliyetle ilişkisi kurulabilir. Varsayımlar açık yazıldığı sürece yaklaşık hesap akademik rapor için yeterlidir.",
            "Rapora eklenecek görsellerin tamamı numaralandırılmış olmalıdır. Her görsel altında kısa ama açıklayıcı başlık bulunmalıdır. Ekran görüntüsü eklenirken yalnızca uygulamanın güzel görünen hali değil, işlemin teknik kanıtı da gösterilmelidir. Örneğin Export Result ekran görüntüsünde sadece liste değil, artifact türleri ve kalite durumu görünmelidir.",
            "Tablolarda 'Güncellenecek' ifadesi final teslimden önce gerçek sayfa numarası, gerçek test sonucu veya gerçek metrikle değiştirilmelidir. Özellikle quality score, upload süresi, video boyutu ve worker süresi gibi alanlar tahmini bırakılmamalıdır. Eğer gerçek WHAM runtime bazı makinelerde çalıştırılamadıysa bu durum dürüstçe 'runtime bağımlılığı nedeniyle fixture/smoke test ile doğrulandı' biçiminde yazılmalıdır.",
        ],
        body_style,
    )
    add_caption(document, "Tablo A.2. Değişiklik kayıt şablonu", caption_style)
    add_table(
        document,
        ["Tarih", "Değişiklik", "Etkilenen alan", "Test kanıtı", "Durum"],
        [
            ["Güncellenecek", "Capture metadata alanı genişletildi.", "Mobil, backend validator, worker input.", "Metadata schema testi, solo upload testi.", "Onaylandı/Bekliyor"],
            ["Güncellenecek", "BVH export doğrulaması eklendi.", "Worker, export_files, ExportResult.", "Blender import smoke testi.", "Onaylandı/Bekliyor"],
            ["Güncellenecek", "Signed URL retry davranışı iyileştirildi.", "Upload manager, backend upload service.", "Bağlantı kesme ve tekrar deneme testi.", "Onaylandı/Bekliyor"],
        ],
        [1.0, 1.8, 2.0, 2.3, 1.2],
    )
    add_caption(document, "Tablo A.3. Ticarileşme maliyet varsayımı şablonu", caption_style)
    add_table(
        document,
        ["Maliyet kalemi", "Varsayım", "Hesap yöntemi"],
        [
            ["GPU işlem", "Ortalama job süresi ve GPU dakika ücreti.", "jobSuresiDakika x gpuDakikaMaliyeti."],
            ["Depolama", "Kaynak video ve artifact toplam GB/ay.", "ortalamaDosyaBoyutu x saklamaSuresi x kullanıcıSayısı."],
            ["Bandwidth", "Export indirme ve video yükleme trafiği.", "GB transfer x sağlayıcı birim fiyatı."],
            ["Destek", "Başarısız job ve kullanıcı soruları.", "aylık destek saati x saatlik operasyon maliyeti."],
            ["Uyum/hukuk", "KVKK metinleri, gizlilik politikası ve veri silme süreci.", "kurulum maliyeti + periyodik danışmanlık."],
        ],
        [1.5, 2.6, 2.7],
    )


def add_ui_section(
    document: Document,
    h1_style: str,
    h2_style: str,
    h3_style: str,
    first_style: str,
    body_style: str,
    caption_style: str,
) -> None:
    add_heading(document, "KULLANICI ARAYÜZÜ", h1_style)
    add_paragraph(
        document,
        "Mocap uygulamasının kullanıcı arayüzü, video tabanlı hareket yakalama sürecinin tamamını adım adım yöneten React Native tabanlı bir mobil uygulamadır. Kullanıcı uygulamayı açtığında proje seçer veya yeni proje oluşturur; ardından çekim modu seçerek kamera kaydına başlar. Kayıt tamamlandıktan sonra take inceleme, upload, processing takibi ve export sonucu ekranları sırayla kullanıcıya çekim sonucunu sunar. Bu bölümde uygulamanın temel ekranları ve kullanıcı akışı ayrıntılandırılmaktadır.",
        first_style,
        first=True,
    )
    add_diagram_box(
        document,
        [
            "Projects ──► Take Listesi ──► Capture Ekrani",
            "                                    |",
            "                             Review Ekrani",
            "                                    |",
            "                          Upload Ilerleme Ekrani",
            "                                    |",
            "                       Processing Status Ekrani",
            "                                    |",
            "                         Export Sonuc Ekrani",
        ],
    )
    add_caption(document, "Şekil 4.1. Mobil uygulama ekran akışı ve navigasyon yapısı", caption_style)

    add_heading(document, "Proje ve Take Yönetim Ekranları", h2_style)
    for text in [
        "Uygulama açıldığında kullanıcı Projects ekranıyla karşılaşır. Bu ekranda kullanıcının oluşturduğu projeler listelenir; her proje bir veya birden fazla take içerebilir. Proje oluşturma akışında proje adı ve opsiyonel açıklama girilir, backend'e POST /projects isteği gönderilir ve sonuç yerel depoya da yazılır. Böylece internet bağlantısı olmasa bile mevcut projeler görüntülenebilir.",
        "Take listesi ekranında seçili projenin tüm take'leri gösterilir. Her take kartında çekim modu (solo, dual, pro), yerel durum (kaydedildi, yüklendi, işleniyor, tamamlandı), oluşturma tarihi ve kısa açıklama yer alır. Kullanıcı yeni take başlatabilir ya da mevcut take'in export sonucuna veya processing durumuna gidebilir. Bu ekran, kullanıcının tüm çekim geçmişini tek noktadan yönetmesini sağlar.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 4.1. Proje ve take yönetimi ekran özellikleri", caption_style)
    add_table(
        document,
        ["Ekran", "Temel görev", "Veri kaynağı"],
        [
            ["Projects listesi", "Projeleri listeler ve yeni proje oluşturur.", "Backend GET /projects + yerel proje deposu."],
            ["Take listesi", "Projeye ait tüm take'leri durum bilgisiyle gösterir.", "Backend GET /projects/:id/takes + yerel take deposu."],
            ["Yeni take başlatma", "Çekim modu ve konfigürasyon seçimi.", "Yerel take oluşturma, backend POST /takes."],
        ],
        [1.8, 2.5, 2.5],
    )

    add_heading(document, "Çekim Ekranı ve Native Kamera Entegrasyonu", h2_style)
    for text in [
        "Capture ekranı uygulamanın teknik merkezini oluşturur. iOS'ta AVFoundation, Android'de CameraX tabanlı native kamera modülü üzerinden çalışır. Ekran yüklendiğinde kamera izni kontrol edilir; izin verilmişse canlı preview başlar. Kullanıcı kayıt düğmesine bastığında video kaydı native modül tarafından başlatılır ve tamamlandığında orijinal video dosyası yerel depoya yazılır.",
        "Capture sırasında aynı zamanda CaptureMetadata paketi üretilir. Bu paket çekim modu, cihaz id, cihaz indeksi, kamera rolü, timestamp, video süresi, çözünürlük, FPS, yön, quality flags ve sync bilgilerini içerir. Metadata paketi video dosyasıyla birlikte yerel take kaydına işlenir ve upload aşamasında ayrı bir nesne olarak backend'e gönderilir. Metadata ve videonun birlikte yönetilmesi, worker tarafında doğru model parametrelerinin kullanılmasını sağlar.",
        "Solo modda tek cihaz kayıt yapar. Dual modda host cihaz oturumu başlatır, guest cihaz join token ile oturuma katılır; her iki cihaz da farklı açılardan kayıt yapar. Pro modda dört cihaz front/right/back/left rolüyle kayıt yapar. Ekranda aktif cihazların bağlantı durumu, kayıt süresi ve mod bilgisi gösterilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 4.2. Çekim modları ve kamera konfigürasyonu", caption_style)
    add_table(
        document,
        ["Mod", "Cihaz sayısı", "Kamera rolü", "Backend sözleşmesi"],
        [
            ["Solo", "1", "primary", "captureMode=solo, expectedVideoCount=1"],
            ["Dual", "2", "front + side", "captureMode=dual, expectedVideoCount=2, host/guest rolleri"],
            ["Pro", "4", "front/right/back/left", "captureMode=pro, expectedVideoCount=4, device index ve rol"],
        ],
        [1.0, 1.2, 1.5, 3.1],
    )

    add_heading(document, "Tekrar İnceleme (Review) ve Metadata Özeti Ekranı", h2_style)
    for text in [
        "Kayıt tamamlandıktan sonra kullanıcı Review ekranına yönlendirilir. Bu ekranda yerel video önizlemesi, çekim süresi, dosya boyutu, FPS, çözünürlük ve kullanılan capture modu görüntülenir. Kullanıcı bu noktada çekimi yeniden yapabilir veya upload aşamasına geçebilir. Review ekranı, yanlış çekimlerin gereksiz yere backend'e gönderilmesini önleyen kalite kapısı işlevi görür.",
        "Ekranın alt kısmında opsiyonel bir metadata özeti paneli bulunur. Bu panel capture metadata paketinin kullanıcı dostu bir özetini gösterir: hangi modda çekildiği, cihaz rolü, FPS ve kalite flag'leri. Teknik kullanıcılar bu bilgilerden çekim kalitesini değerlendirerek ilerleyip ilerlememek konusunda daha bilinçli karar verebilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Upload İlerleme Ekranı", h2_style)
    for text in [
        "Upload ekranı, metadata ve video dosyasının backend aracılığıyla S3 uyumlu nesne depolamaya güvenli bir şekilde aktarılmasını yönetir. Uygulama önce backend'den upload session ve imzalı URL'ler alır. Metadata JSON ayrı bir object key'e, video dosyası ayrı bir object key'e PUT isteğiyle yüklenir. Yükleme tamamlandığında backend'e complete çağrısı yapılır; backend object storage üzerinde nesne varlığını ve boyutu doğrular.",
        "Upload ilerleme ekranında yüzde bazlı ilerleme çubuğu, aktif adım bilgisi (metadata upload, video upload, tamamlama doğrulama) ve hata durumunda yeniden deneme seçeneği yer alır. Bağlantı kesintisi veya signed URL süresi dolması durumunda uygulama fresh URL alarak upload'ı kaldığı yerden yeniden başlatabilir. Bu yaklaşım büyük video dosyaları için güvenilir ve kullanıcı dostu bir deneyim sağlar.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "POST /uploads/init -> signed PUT URL'leri al",
            "PUT metadata.json  -> S3 object key'e",
            "PUT source_video   -> S3 object key'e",
            "POST /uploads/complete -> object size dogrulama",
        ],
    )
    add_caption(document, "Şekil 4.2. Upload akışı ve signed URL yönetimi", caption_style)

    add_heading(document, "İşlem Durumu (Processing Status) Ekranı", h2_style)
    for text in [
        "Processing Status ekranı kullanıcıya backend worker'ın çekimi işleme durumunu gerçek zamanlı gösterir. Ekran yüklendiğinde backend'e GET /process/:jobId isteği gönderilir ve job state okunur. Mobil uygulama belirli aralıklarla bu endpoint'i sorgulayarak job'ın mevcut aşamasını (queued, ingesting, extracting_frames, solving_motion, cleaning, exporting, succeeded, failed) kullanıcıya aktarır.",
        "Her aşama için açıklayıcı bir metin gösterilir. Örneğin 'solving_motion' aşamasında 'Hareket verisi analiz ediliyor' mesajı görünür. Hata durumunda hata kodu ve anlaşılır bir açıklama sunulur; kullanıcı gerekirse işlemi iptal edebilir. Başarı durumunda kullanıcı otomatik olarak Export Sonuç ekranına yönlendirilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "queued -> ingesting -> extracting_frames -> solving_motion",
            "                                                 |",
            "                                             cleaning",
            "                                                 |",
            "                           exporting -> succeeded -> ExportResult",
            "                               |",
            "                         failed / canceled",
        ],
    )
    add_caption(document, "Şekil 4.3. Processing job durum akışı ve ekran geçişleri", caption_style)

    add_heading(document, "Export Sonuç Ekranı: BVH, Kalite Raporu ve Önizleme", h2_style)
    for text in [
        "Export Sonuç ekranı, worker tarafından üretilen artifact dosyalarını kullanıcıya listeler. Her artifact için dosya adı, format, boyut ve indirme seçeneği gösterilir. Kullanıcı BVH dosyasını indirerek doğrudan Blender, Maya, MotionBuilder, Unity veya Unreal Engine'e aktarabilir. Solved motion JSON, SMPL parameters ve kalite raporu teknik kullanıcılar için erişilebilir durumdadır.",
        "Kalite raporu özeti ekranın üst kısmında gösterilir. Score, grade, uyarı sayısı ve varsa kritik hatalar burada özetlenir. Kullanıcı kalite raporunun detayına giderek hangi aşamada ne tür bir sorun oluştuğunu inceleyebilir. Çoklu kamera denemelerinde multi-view diagnostics bölümü de gösterilir; reconstruction availability, sync confidence ve finalAnimationSource gibi bilgiler sunulur.",
        "Overlay preview videosu varsa ekranda oynatılabilir. Bu video orijinal kaynak üzerine WHAM hareket çözümünün görsel olarak bindirildiği bir animasyon önizlemesidir. Kullanıcı hem orijinal videoyu hem de çözüm sonucunu karşılaştırma imkânı bulur.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 4.3. Export sonuç ekranındaki artifact listesi", caption_style)
    add_table(
        document,
        ["Artifact", "Gösterim", "Kullanım amacı"],
        [
            ["animation.bvh", "İndirilebilir dosya", "Blender, Maya, Unity, Unreal ve diğer DCC araçlarına aktarım."],
            ["solved_motion.json", "İndirilebilir dosya", "Eklem pozisyonları, root motion ve kare bazlı hareket verisi."],
            ["quality_report.json", "Ekranda özet + detay", "Teknik kalite skoru, uyarılar ve validasyon sonucu."],
            ["preview_overlay.mp4", "Ekranda oynatılabilir", "Hareket çözümünün görsel kontrolü için önizleme."],
            ["smpl_parameters.json", "İndirilebilir dosya", "SMPL pose, shape ve kamera/model parametreleri."],
            ["motion_pipeline_report.json", "İndirilebilir dosya", "Stage süreleri, artifact lineage ve finalAnimationSource."],
        ],
        [2.0, 1.6, 3.2],
    )

    add_heading(document, "Çoklu Kamera Kurulum ve Multi-View Diagnostik Ekranı", h2_style)
    for text in [
        "Dual ve pro çoklu kamera modlarında kullanıcı önce MultiViewSetup ekranına yönlendirilir. Bu ekranda host cihaz bir capture session oluşturur ve join token üretir. Guest cihazlar bu token'ı girerek oturuma katılır. Her cihazın bağlantı durumu, atanan kamera rolü ve device index bilgisi ekranda gösterilir. Tüm beklenen cihazlar bağlandığında kayıt başlatma düğmesi aktif hale gelir.",
        "Kayıt sonrasında Processing Status ve Export Sonuç ekranlarında multi-view diagnostics bölümü görünür. Bu bölümde matched frame count, sync confidence, reconstruction availability ve finalAnimationSource alanları gösterilir. Kalibrasyon eksikse veya senkronizasyon başarısız olduysa sistem bunu açıkça belirtir; kullanıcıya sahte bir 'çoklu kamera başarısı' sunulmaz.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 4.4. Multi-view kurulum ve diagnostik ekran bilgileri", caption_style)
    add_table(
        document,
        ["Bilgi alanı", "Kaynak", "Gösterim yeri"],
        [
            ["Join token", "Backend capture session", "MultiViewSetup ekranında host/guest paylaşımı."],
            ["Cihaz rolü ve indeksi", "Backend device role ataması", "Her cihaz kartında rol ve indeks etiketi."],
            ["Bağlantı durumu", "WebSocket relay / backend status", "Gerçek zamanlı bağlantı göstergesi."],
            ["Sync confidence", "Worker multi_view_sync.json", "Export sonuç ekranındaki diagnostics bölümü."],
            ["Reconstruction availability", "quality_report.multiView", "Export ekranında uyarı veya bilgi mesajı."],
            ["Final animation source", "quality_report.finalAnimationSource", "Export ekranında 'primary_wham' veya 'true_dual_solve'."],
        ],
        [2.0, 2.3, 2.5],
    )

    add_heading(document, "Backend API Entegrasyonu ve Uygulama Durum Yönetimi", h2_style)
    for text in [
        "Mobil uygulama backend ile MocapApiClient sınıfı üzerinden haberleşir. Bu sınıf tüm HTTP isteklerini yönetir, token header'ı ekler, hata yanıtlarını yapılandırılmış ApiError'a dönüştürür ve retry kabiliyeti sunar. Servis katmanı API sonuçlarını domain modellerine dönüştürür; bu sayede ekranlar doğrudan ham JSON ile değil, type-safe domain nesneleriyle çalışır.",
        "Uygulama durum yönetiminde Zustand tabanlı captureStore ve multiViewStore kullanılmaktadır. captureStore aktif çekim oturumunu, kayıt durumunu, yerel take meta'sını ve upload/processing state'ini yönetir. multiViewStore çoklu kamera oturumundaki cihaz listesini, bağlantı durumlarını, sync state'ini ve WebSocket relay haberleşmesini koordine eder.",
        "Yerel take deposu (TakeRepository) cihaz dosya sisteminde meta.json ve JSONL chunk dosyaları kullanarak kayıtları saklar. Bu yapı sayesinde internet bağlantısı olmasa dahi tamamlanan çekimler kaybolmaz; uygulama yeniden açıldığında upload edilmemiş take'ler tespit edilir ve kullanıcıya upload etmesi önerilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 4.5. Mobil uygulama yazılım katmanları", caption_style)
    add_table(
        document,
        ["Katman", "Temel sınıf/modül", "Görev"],
        [
            ["Ekranlar", "CaptureScreen, ReviewScreen, UploadScreen, ProcessingScreen, ExportScreen", "Kullanıcı etkileşimi ve görsel sunum."],
            ["State / Hook", "captureStore, multiViewStore, useRecorder, useWhamCapture", "Canlı durum yönetimi ve native kamera koordinasyonu."],
            ["Servis", "MocapApiClient, SignedUrlUploadManager", "Backend haberleşmesi ve upload yönetimi."],
            ["Domain", "Take, CaptureMetadata, PoseFrame, ExportResult", "Projenin temel veri modelleri."],
            ["Yerel depo", "TakeRepository, LocalFileStore", "Cihaz üzerindeki take ve video dosyası yönetimi."],
        ],
        [1.4, 3.0, 2.4],
    )


def add_results_section(
    document: Document,
    h1_style: str,
    h2_style: str,
    first_style: str,
    body_style: str,
    caption_style: str,
) -> None:
    add_heading(document, "TEST VE SONUÇLAR", h1_style)
    add_paragraph(
        document,
        "Bu bölümde Mocap sisteminin test yöntemi, kabul kriterleri, gerçekleştirilen testlerin sonuçları ve teknik bulgular sunulmaktadır. Test süreci üç seviyede yürütülmüştür: mobil kullanıcı akışı testleri, backend/worker doğrulama testleri ve artifact kalite testleri. Her test için giriş verisi, beklenen sonuç ve gözlenen sonuç ayrı ayrı belirtilmiştir.",
        first_style,
        first=True,
    )

    add_heading(document, "Uçtan Uca Test Yöntemi ve Kabul Kriterleri", h2_style)
    for text in [
        "Test süreci önce sanal laboratuvar ortamında (yerel backend, MinIO ve fixture veri) başlar; ardından gerçek mobil cihazlarla ve gerçek WHAM runtime ile doğrulanır. Sanal laboratuvar testleri metadata doğrulaması, upload akışı, job state geçişleri ve export schema kontrollerini kapsar. Gerçek cihaz testleri ise solo, dual ve opsiyonel pro kamera senaryolarında uçtan uca çekim-upload-processing-export döngüsünü doğrular.",
        "Temel kabul kriterleri şunlardır: video dosyası oluşturulur ve metadata geçerlidir; metadata ve video signed URL ile yüklenir; upload complete backend doğrulamasını geçer; processing job terminal state'e (succeeded/failed) ulaşır; BVH ve JSON artifact dosyaları export_files tablosuna kaydedilir; quality_report okunabilir score/grade/warnings içerir; BVH dosyası Blender'da yapısal import hatası vermez.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 6.1. Uçtan uca test aşaması ve kabul kriterleri", caption_style)
    add_table(
        document,
        ["Test aşaması", "Başarı kriteri", "Kanıt türü"],
        [
            ["Kamera kaydı", "Video dosyası ve metadata oluşur.", "Video metadata tablosu, Capture ekranı görüntüsü."],
            ["Signed URL upload", "Metadata ve video S3'e yüklenir.", "Upload progress ekranı, backend complete kaydı."],
            ["Processing job", "Job terminal state'e ulaşır.", "Job timeline, state geçişleri."],
            ["Artifact üretimi", "BVH ve JSON raporlar export_files'a kaydedilir.", "Export listesi ve dosya boyutları."],
            ["Kalite raporu", "Score/grade/warnings okunabilir.", "quality_report özet tablosu."],
            ["BVH import", "Blender yapısal import hatası vermez.", "Blender smoke test veya ekran görüntüsü."],
            ["Multi-view diagnostic", "Eksik kalibrasyonda sahte başarı üretilmez.", "diagnostic_only veya missing_calibration uyarısı."],
        ],
        [1.8, 2.4, 2.6],
    )

    add_heading(document, "Performans ve Kalite Testleri", h2_style)
    for text in [
        "Performans testleri upload süresi, video normalize süresi, WHAM/SMPL solve süresi ve toplam pipeline süresi metriklerini kapsar. Bu metrikler motion_pipeline_report_json içindeki stage sürelerinden okunur. Geliştirme ortamında ortalama 10 saniyelik bir solo video için normalize süresi birkaç saniye, WHAM solve süresi ise GPU kapasitesine bağlı olarak değişmektedir.",
        "Kalite testleri quality_report_json içindeki metrikler üzerinden yapılır. Değerlendirilen metrikler arasında frame count ve video süresi tutarlılığı, root motion stabilitesi, eklem rotasyonlarında NaN/Infinity bulunmaması ve BVH import başarısı yer alır. Çoklu kamera testlerinde ek olarak matched frame count, sync confidence ve reconstruction availability metrikleri değerlendirilir.",
        "Düşük ışık, hızlı hareket, kısmi vücut görünürlüğü ve bozuk metadata gibi sınır durum testleri de yapılmıştır. Bu testlerde quality_report'un anlamlı uyarı üretebildiği, worker'ın sahte başarı sunmak yerine açıklayıcı hata kodu döndürebildiği doğrulanmıştır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 6.2. Test senaryoları ve kalite sonuçları", caption_style)
    add_table(
        document,
        ["Senaryo", "Koşul", "Gözlenen sonuç"],
        [
            ["Solo iyi koşul", "Yeterli ışık, tüm vücut görünür, 10-20 sn.", "Job succeeded, BVH export, quality grade yüksek."],
            ["Solo düşük ışık", "Yetersiz ortam aydınlatması.", "quality_report warnings, düşük confidence notu."],
            ["Solo hızlı hareket", "Ani yön değişimi ve hızlı hareket.", "cleanup_report jitter/velocity uyarıları."],
            ["Eksik kadraj", "Bacaklar veya üst vücut kısmen görünmez.", "Tracking warning, quality score düşüşü."],
            ["Dual kamera", "İki cihaz, host/guest rolü.", "Matched frame count, sync confidence, diagnostic raporu."],
            ["Bozuk metadata", "Eksik captureMode veya hatalı schema.", "Job başlatılmadan açıklayıcı hata mesajı."],
        ],
        [1.8, 2.3, 2.7],
    )

    add_heading(document, "Test Sonuçları ve Teknik Bulgular", h2_style)
    for text in [
        "Bu çalışma sonucunda Mocap projesi için mobil cihaz odaklı, video tabanlı ve backend destekli bir işaretleyicisiz hareket yakalama mimarisi ortaya konmuştur. Sistem mobil uygulama ile orijinal video kaydını alabilmekte, capture metadata üretebilmekte, yerel take kaydı tutabilmekte, backend'e signed URL üzerinden upload yapabilmekte, processing job oluşturabilmekte ve worker tarafından üretilen export dosyalarını kullanıcıya sunabilmektedir.",
        "Projenin en güçlü yönü, hareket yakalama işlemini tek bir cihaz üstü anlık tahmin olarak ele almaması; orijinal videoyu ana üretim kaynağı kabul ederek backend'de tekrar işlenebilir bir pipeline kurmasıdır. Bu yaklaşım model güncellemesi, kalite raporu üretimi, GPU kullanımı ve farklı export formatlarının desteklenmesi açısından sürdürülebilir bir mimari sağlar.",
        "Backend'in imzalı URL mimarisi, büyük video dosyalarını API sunucusu üzerinden geçirmeden nesne depolamaya aktarmasına olanak tanımakta; API sunucusunun bant genişliği yükünü anlamlı biçimde azaltmaktadır. Worker pipeline'ının modüler aşama yapısı (ingest → normalize → solve → cleanup → export → report), her aşamanın bağımsız izlenmesini ve hata durumunda anlamlı geri bildirim üretilmesini sağlamaktadır.",
        "Dual ve pro çoklu kamera altyapısının oluşturulması önemli bir mühendislik çıktısıdır. Session yönetimi, cihaz rolü, join token, frame sync, kalibrasyon ve triangulation için altyapı hazırlanmış; multi-view diagnostic raporlar tanısal kalite metrikleri sunmaktadır. Final animasyonun mevcut durumda primary WHAM solve üzerinden üretilmesi açık bir teknik sınır olarak belgelenmiştir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 6.3. Mimari kazanımlar ve sistem çıktıları", caption_style)
    add_table(
        document,
        ["Bileşen / Özellik", "Gerçekleştirilen çıktı", "Değeri"],
        [
            ["Mobil kayıt", "Native kamera, capture metadata, yerel take deposu.", "Stüdyo kurmadan video kaydı ve metadata üretimi."],
            ["Upload mimarisi", "Signed URL, object size doğrulama, retry.", "Büyük video aktarımında güvenilir ve ölçeklenebilir yol."],
            ["Worker pipeline", "Normalize, WHAM solve, cleanup, BVH, kalite raporu.", "Modelden kalite raporuna tam üretim hattı."],
            ["Export seti", "BVH, solved_motion, quality_report, pipeline_report.", "DCC uyumlu çıktı ve teknik doğrulama artifact'leri."],
            ["Multi-view altyapısı", "Session, sync, calibration, triangulation, diagnostics.", "Gelecekte true multi-view solve için hazır temel."],
            ["Kalite raporu", "Score, grade, warnings, finalAnimationSource.", "Sonucun sınırlarını açıklayan şeffaf çıktı."],
        ],
        [2.0, 2.5, 2.3],
    )

    add_heading(document, "Gelecek Çalışmalar İçin Öneriler", h2_style)
    for text in [
        "Birinci öneri gerçek üretim seviyesinde kimlik doğrulama ve yetkilendirme eklenmesidir. Mevcut bearer token yaklaşımı prototip için yeterlidir; ancak çok kullanıcılı bir servis için JWT/OAuth tabanlı auth, token yenileme ve kullanıcı-take sahipliği kontrolü eklenmelidir.",
        "İkinci öneri gerçek cihazlardan oluşan golden fixture seti hazırlamaktır. Solo, dual ve pro çekimlerle kapsamlı golden fixture hazırlandığında her pipeline değişikliğinden sonra karşılaştırmalı kalite testi mümkün hale gelir.",
        "Üçüncü öneri audio/native sync, AprilTag/checkerboard calibration ve triangulated 3D constraints'in WHAM/SMPL solve'a entegre edilmesidir. Bu geliştirme tamamlandığında dual/pro kamera akışından gerçek anlamda iyileştirilmiş final animasyon üretilebilecektir.",
        "Dördüncü öneri Blender/Unity/Unreal import testlerinin otomatik QA sürecine alınmasıdır. BVH dosyasının DCC aracı tarafından import edilebilmesi, export kalitesinin en somut kanıtıdır; bu testin otomatikleştirilmesi pipeline'ın güvenilirliğini artırır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "Bugun:  Solo WHAM final animation + multi-view diagnostics",
            "Kisa:   Real-device QA, auth hardening, golden fixture seti",
            "Orta:   Calibration clip, audio/native sync, model preflight",
            "Uzun:   True multi-view solve, production deploy, privacy policy",
        ],
    )
    add_caption(document, "Şekil 6.1. Test sonuçları ve gelecek geliştirme fazları", caption_style)


def add_project_plan_section(
    document: Document,
    h1_style: str,
    h2_style: str,
    first_style: str,
    body_style: str,
    caption_style: str,
) -> None:
    add_heading(document, "PROJE PLANI VE MALİYET ANALİZİ", h1_style)
    add_paragraph(
        document,
        "Bu bölümde Mocap projesinin geliştirme sürecinde izlenen iş-zaman planı ve projenin tahmini maliyet analizi sunulmaktadır. İş-zaman çizelgesi analiz, tasarım, geliştirme, test ve raporlama aşamalarını kapsamakta; maliyet analizi ise donanım, yazılım, cloud servisleri ve geliştirme maliyetlerini ayrıntılandırmaktadır.",
        first_style,
        first=True,
    )

    add_heading(document, "İş-Zaman Çizelgesi", h2_style)
    add_paragraphs(
        document,
        [
            "Proje geliştirme süreci beş ana aşamada planlanmıştır. İlk aşamada (analiz ve mimari tasarım) hareket yakalama problemi incelenmiş, mobil uygulama, backend API, nesne depolama ve worker pipeline gereksinimleri belirlenmiş; veri modeli ve API sözleşmeleri tasarlanmıştır.",
            "İkinci aşamada (mobil uygulama geliştirme) React Native/Expo Dev Client tabanlı mobil uygulama geliştirilmiş; kamera kayıt modülü, take deposu, upload yöneticisi ve tüm kullanıcı akışı ekranları tamamlanmıştır. Üçüncü aşamada (backend ve worker geliştirme) Fastify backend API servisleri, PostgreSQL veritabanı migrasyonları, S3 upload/download akışları, processing job kuyruğu ve worker pipeline hazırlanmıştır.",
            "Dördüncü aşama (test ve kalite güvence) backend birim testleri, triangulation/frame-sync/calibration testleri, synthetic/golden QA testleri ve gerçek cihaz testlerini kapsamıştır. Beşinci ve son aşama (raporlama ve sunum hazırlığı) tüm belgelerin hazırlanması, ekran görüntüleri ve kanıt dosyalarının eklenmesi ve sözlü sınav sunumunun oluşturulmasını içermektedir.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 7.1. Proje iş-zaman çizelgesi", caption_style)
    add_table(
        document,
        ["Aşama", "Kapsam", "Tahmini süre"],
        [
            ["Analiz ve mimari tasarım", "Gereksinim analizi, veri modeli, API sözleşmeleri.", "3–4 hafta"],
            ["Mobil uygulama geliştirme", "Kamera modülü, take deposu, upload, tüm ekranlar.", "5–6 hafta"],
            ["Backend ve worker geliştirme", "API servisleri, PostgreSQL, S3 akışı, worker pipeline.", "6–7 hafta"],
            ["Test ve kalite güvence", "Birim testler, QA testleri, gerçek cihaz testleri.", "3–4 hafta"],
            ["Raporlama ve sunum", "Tez yazımı, kanıt toplama, sözlü sınav hazırlığı.", "2–3 hafta"],
        ],
        [2.0, 3.0, 1.8],
    )

    add_heading(document, "Maliyet Analizi", h2_style)
    add_paragraphs(
        document,
        [
            "Proje maliyet analizi donanım, yazılım/lisans, bulut servisleri ve geliştirme maliyetleri olmak üzere dört kategoride değerlendirilmiştir. Donanım kategorisinde geliştirme bilgisayarları ve test için kullanılan mobil cihazlar (iOS ve Android) yer almaktadır.",
            "Yazılım ve lisans kategorisinde React Native, Fastify, PostgreSQL, MinIO gibi açık kaynak bileşenler maliyetsiz kullanılmıştır. WHAM modeli araştırma lisansıyla, SMPL modeli ise kendi lisans koşullarıyla kullanılmakta olup ticari kullanım için ayrı lisans anlaşması gerekecektir.",
            "Bulut servisleri kategorisinde en önemli maliyet GPU işleme (RunPod Serverless), nesne depolama (S3 uyumlu) ve bant genişliğidir. Geliştirme ve test aşamasında video başına ortalama GPU maliyeti, video süresi ve model inference süresine bağlı olarak değişmektedir.",
        ],
        body_style,
    )
    add_caption(document, "Tablo 7.2. Proje maliyet analizi", caption_style)
    add_table(
        document,
        ["Maliyet kategorisi", "Kalem", "Durum / Tahmini Tutar"],
        [
            ["Donanım", "Geliştirici bilgisayarları", "Mevcut ekipman (ekstra maliyet yok)"],
            ["Donanım", "iOS + Android test cihazları", "Mevcut cihazlar (ekstra maliyet yok)"],
            ["Yazılım / Lisans", "React Native, Fastify, PostgreSQL, MinIO", "Açık kaynak (ücretsiz)"],
            ["Yazılım / Lisans", "WHAM araştırma lisansı", "Akademik kullanım için uygun; ticari lisans gerekebilir"],
            ["Yazılım / Lisans", "SMPL model lisansı", "Araştırma lisanslı; ticari kullanım ayrı anlaşma gerektirir"],
            ["Bulut – GPU", "RunPod Serverless (WHAM solve)", "Video başına ~0.05–0.20 USD (GPU hızına göre değişir)"],
            ["Bulut – Depolama", "S3 uyumlu nesne depolama", "GB başına ~0.01–0.02 USD/ay"],
            ["Bulut – Bant genişliği", "Upload + download trafiği", "GB başına ~0.01–0.09 USD (sağlayıcıya göre değişir)"],
            ["Geliştirme zamanı", "Ekip geliştirme ve test süresi", "Akademik proje kapsamında (fırsat maliyeti)"],
        ],
        [1.8, 2.5, 2.5],
    )


def build_report() -> None:
    document = Document(str(TEMPLATE))
    clear_body(document)

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1)
    add_page_number(section)

    body_style = style_name(document, "Ana_Paragraf_Yazi_Stili_Sau")
    first_style = style_name(document, "Bolum_Ilk_Paragraf_Sau", body_style)
    front_h_style = style_name(document, "Ilk_Sayfalar_Basligi_Sau", "Normal")
    h1_style = style_name(document, "Heading 1")
    h2_style = style_name(document, "Alt_Baslık_Sau", "Heading 2")
    h3_style = style_name(document, "Ikincil_Alt_Baslik_Sau", "Heading 3")
    caption_style = style_name(document, "Caption")

    cover_lines = [
        "T.C.",
        "SAKARYA ÜNİVERSİTESİ",
        "BİLGİSAYAR VE BİLİŞİM BİLİMLERİ FAKÜLTESİ",
        "",
        "BSM 498 BİTİRME ÇALIŞMASI",
        "",
        "MOBİL CİHAZLAR İLE VİDEO TABANLI İŞARETLEYİCİSİZ HAREKET YAKALAMA SİSTEMİ GELİŞTİRMESİ",
        "",
        "B231210351 - BURAK COŞKUN",
        "Öğrenci2 No - BERKE PİTE",
        "B221210059 - EREN OZAN PİTE",
        "",
        "Proje Adı: Mocap",
        "2025-2026 Bahar Dönemi",
    ]
    for line in cover_lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14 if line and line.isupper() else 12)
        r.bold = bool(line and (line.isupper() or line.startswith("B")))
    document.add_page_break()

    add_heading(document, "ÖNSÖZ", front_h_style)
    for text in [
        "Bu bitirme çalışması, mobil cihazlarla kaydedilen videoları kullanarak işaretleyicisiz hareket yakalama sürecini uçtan uca ele alan Mocap projesini açıklamak amacıyla hazırlanmıştır. Proje, yalnızca kamera görüntüsü üzerinden anlık bir görsel efekt üretmeyi değil; çekim, metadata üretimi, güvenli yükleme, bulut tabanlı işleme, kalite değerlendirmesi ve animasyon dışa aktarımı aşamalarını birlikte kapsayan üretime dönük bir sistem mimarisi oluşturmayı hedeflemektedir.",
        "Çalışma boyunca hareket yakalama probleminin mobil cihaz koşullarındaki zorlukları, video tabanlı tek kamera çözümünün avantajları, çoklu kamera altyapısının sağladığı ek tanısal veriler, backend iş kuyruğunun sorumlulukları ve WHAM/SMPL tabanlı model çıktılarının animasyon üretim sürecindeki yeri incelenmiştir. Bu rapor, proje kod tabanı, mimari dokümanları, mobil ekran akışları, backend servisleri, worker işlem hattı ve test/QA planları temel alınarak hazırlanmıştır.",
        "Projenin geliştirilmesi sırasında mobil uygulama, yerel veri saklama, backend API, nesne depolama, iş kuyruğu, video normalizasyonu, hareket çözümü ve kalite raporlaması gibi farklı alanlar bir arada ele alınmıştır. Bu nedenle raporda her bir bileşen ayrı ayrı tanıtılmış, ardından bu bileşenlerin tek bir hareket yakalama ürünü olarak nasıl bütünleştiği açıklanmıştır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    document.add_page_break()

    add_heading(document, "İÇİNDEKİLER", front_h_style)
    toc_rows = [
        ["ÖNSÖZ", "iii"],
        ["İÇİNDEKİLER", "iv"],
        ["SİMGELER VE KISALTMALAR LİSTESİ", "vi"],
        ["ŞEKİLLER LİSTESİ", "vii"],
        ["TABLOLAR LİSTESİ", "viii"],
        ["ÖZET", "ix"],
        ["BÖLÜM 1.  GİRİŞ", "1"],
        ["1.1. Problemin Tanımı: Hareket Yakalamada Erişilebilirlik Sorunu", "1"],
        ["1.2. Projenin Amacı ve Kapsamı", "2"],
        ["1.3. Önerilen Sistemin Ana Hatları", "3"],
        ["BÖLÜM 2.  SİSTEMATİK YAKLAŞIM", "5"],
        ["2.1. Video Tabanlı Hareket Yakalama Formülasyonu", "5"],
        ["2.2. Genel Sistem Mimarisi", "6"],
        ["2.3. Mobil Uygulama ve Donanım Katmanı", "7"],
        ["2.4. Backend API, Veri Modeli ve PostgreSQL Yapısı", "8"],
        ["2.5. Worker Pipeline ve Model Çalışma Ortamı (WHAM/SMPL)", "9"],
        ["2.6. Bulut, S3 Nesne Depolama ve İş Kuyruğu Mimarisi", "10"],
        ["BÖLÜM 3.  DENEY DÜZENEĞİ VE SANAL LABORATUVAR", "12"],
        ["3.1. Uygulama İş Akışı (User Flow) ve Kullanım Senaryoları", "12"],
        ["3.2. Tek ve Çoklu Kamera Deney Düzeneği", "13"],
        ["3.3. Sanal Laboratuvar: Bileşenler ve Yerel Test Ortamı", "14"],
        ["BÖLÜM 4.  KULLANICI ARAYÜZÜ", "16"],
        ["4.1. Proje ve Take Yönetim Ekranları", "16"],
        ["4.2. Çekim Ekranı ve Native Kamera Entegrasyonu", "17"],
        ["4.3. Tekrar İnceleme (Review) ve Metadata Özeti Ekranı", "18"],
        ["4.4. Upload İlerleme Ekranı", "19"],
        ["4.5. İşlem Durumu (Processing Status) Ekranı", "20"],
        ["4.6. Export Sonuç Ekranı: BVH, Kalite Raporu ve Önizleme", "21"],
        ["4.7. Çoklu Kamera Kurulum ve Multi-View Diagnostik Ekranı", "22"],
        ["4.8. Backend API Entegrasyonu ve Uygulama Durum Yönetimi", "23"],
        ["BÖLÜM 5.  VERİ GÜVENLİĞİ DEĞERLENDİRMESİ", "25"],
        ["5.1. Video ve Metadata Bütünlüğü", "25"],
        ["5.2. Güvenli Yükleme ve Signed URL Yönetimi", "26"],
        ["5.3. WHAM/SMPL Model Varlık Güvenliği ve API Güvenliği", "27"],
        ["5.4. Kişisel Verilerin Korunması ve KVKK Uyumu", "28"],
        ["BÖLÜM 6.  TEST VE SONUÇLAR", "30"],
        ["6.1. Uçtan Uca Test Yöntemi ve Kabul Kriterleri", "30"],
        ["6.2. Performans ve Kalite Testleri", "31"],
        ["6.3. Test Sonuçları ve Teknik Bulgular", "32"],
        ["6.4. Gelecek Çalışmalar İçin Öneriler", "33"],
        ["BÖLÜM 7.  PROJE PLANI VE MALİYET ANALİZİ", "35"],
        ["7.1. İş-Zaman Çizelgesi", "35"],
        ["7.2. Maliyet Analizi", "36"],
        ["BÖLÜM 8.  DEĞİŞİKLİK YÖNETİMİ VE MÜHENDİSLİK STANDARTLARI", "37"],
        ["8.1. Mocap Yazılım Geliştirme Süreci", "37"],
        ["8.2. Sürümleme, Etki Analizi ve Değişiklik Kontrolü", "38"],
        ["8.3. Mühendislik Standartları ve Kalite Kapıları", "39"],
        ["8.4. Ekip Sorumlulukları ve Yayınlama Süreci", "40"],
        ["BÖLÜM 9.  TİCARİLEŞME PLANI", "41"],
        ["9.1. Pazar Analizi ve Hedef Kitle", "41"],
        ["9.2. Gelir Modeli ve Maliyet Yapısı", "42"],
        ["9.3. Rekabet Avantajı", "43"],
        ["9.4. Büyüme Stratejisi ve Yol Haritası", "44"],
        ["KAYNAKLAR", "46"],
        ["EK A", "48"],
        ["ÖZGEÇMİŞ", "51"],
        ["BSM 498 BİTİRME ÇALIŞMASI DEĞERLENDİRME VE SÖZLÜ SINAV TUTANAĞI", "52"],
    ]
    add_table(document, ["Başlık", "Sayfa"], toc_rows, [5.6, 0.8])
    document.add_page_break()

    add_heading(document, "SİMGELER VE KISALTMALAR LİSTESİ", front_h_style)
    add_table(
        document,
        ["Kısaltma", "Açıklama"],
        [
            ["API", "Uygulama programlama arayüzü; mobil uygulama ile backend arasındaki sözleşmeli HTTP uçları."],
            ["BVH", "Biovision Hierarchy; iskelet hiyerarşisi ve hareket kareleri içeren animasyon aktarım formatı."],
            ["CI/CD", "Sürekli entegrasyon ve sürekli teslim; değişikliklerin test kapılarından geçerek yayınlanması yaklaşımı."],
            ["DCC", "Digital Content Creation; Blender, Maya, MotionBuilder gibi içerik üretim araçları."],
            ["DLT", "Direct Linear Transform; iki veya daha fazla kamera görüntüsünden 3B nokta kestirimi için kullanılan yaklaşım."],
            ["FPS", "Saniyedeki kare sayısı; çekim, önizleme ve export kalitesini etkileyen temel ölçüt."],
            ["GPU", "Grafik işlem birimi; WHAM/SMPL gibi ağır model çıkarımlarında kullanılan hızlandırıcı donanım."],
            ["JSON/JSONL", "Yapılandırılmış veri formatları; metadata, kalite raporu ve yerel frame chunk verileri için kullanılır."],
            ["KVKK", "Kişisel Verilerin Korunması Kanunu; kullanıcı videoları ve hareket verileri için dikkate alınması gereken hukuki çerçeve."],
            ["MVP", "Minimum uygulanabilir ürün; ticarileşme planında ilk pazara çıkış sürümü."],
            ["ROI", "Yatırım getirisi; ticarileşme ve maliyet analizinde kullanılan oran."],
            ["S3", "Nesne depolama API modeli; video ve artifact dosyalarının saklanmasında kullanılır."],
            ["SLA", "Hizmet seviyesi anlaşması; ticari kullanımda işlem süresi ve erişilebilirlik hedeflerini ifade eder."],
            ["SMPL", "Parametrik insan vücut modeli; vücut pozu, şekli ve eklem konumlarını temsil eder."],
            ["SMPLify", "Görüntüden elde edilen vücut parametrelerini optimize ederek SMPL uyumunu iyileştiren süreç."],
            ["WHAM", "World-grounded Humans with Accurate Motion; videodan dünya koordinatlarında insan hareketi çıkaran model hattı."],
            ["QA", "Quality Assurance; çekimden export dosyasına kadar doğrulama ve kalite güvence süreci."],
            ["WebSocket", "Çift yönlü gerçek zamanlı haberleşme protokolü; çoklu kamera eşleşme/relay akışında kullanılır."],
        ],
        [1.2, 5.5],
    )
    document.add_page_break()

    add_heading(document, "ŞEKİLLER LİSTESİ", front_h_style)
    add_table(
        document,
        ["Şekil", "Açıklama", "Sayfa"],
        [
            ["Şekil 1.1.", "Mocap genel iş akışı", "Güncellenecek"],
            ["Şekil 2.1.", "Sistem mimarisi blok diyagramı", "Güncellenecek"],
            ["Şekil 2.2.", "Mobil, backend ve worker veri akışı", "Güncellenecek"],
            ["Şekil 2.3.", "Bulut işleme ve artifact üretim hattı", "Güncellenecek"],
            ["Şekil 3.1.", "Temel kullanım senaryosu iş akışı", "Güncellenecek"],
            ["Şekil 3.2.", "Tek kamera ve çoklu kamera deney düzeneği", "Güncellenecek"],
            ["Şekil 3.3.", "Processing job durum akışı", "Güncellenecek"],
            ["Şekil 4.1.", "Mobil uygulama ekran akışı ve navigasyon yapısı", "Güncellenecek"],
            ["Şekil 4.2.", "Upload akışı ve signed URL yönetimi", "Güncellenecek"],
            ["Şekil 4.3.", "Processing job durum akışı ve ekran geçişleri", "Güncellenecek"],
            ["Şekil 5.1.", "Veri güvenliği ve bütünlük kontrol akışı", "Güncellenecek"],
            ["Şekil 6.1.", "Test sonuçları ve gelecek geliştirme fazları", "Güncellenecek"],
            ["Şekil 8.1.", "Değişiklik yönetimi akışı", "Güncellenecek"],
            ["Şekil 8.2.", "Release kapısı ve kalite güvence zinciri", "Güncellenecek"],
            ["Şekil 9.1.", "Ticarileşme yol haritası", "Güncellenecek"],
        ],
        [1.0, 4.8, 1.2],
    )
    document.add_page_break()

    add_heading(document, "TABLOLAR LİSTESİ", front_h_style)
    add_table(
        document,
        ["Tablo", "Açıklama", "Sayfa"],
        [
            ["Tablo 1.1.", "Problem, amaç ve kapsam matrisi", "Güncellenecek"],
            ["Tablo 1.2.", "Hedef kullanıcı ve değer önerisi matrisi", "Güncellenecek"],
            ["Tablo 1.3.", "Proje kapsamı ve kapsam dışı kararlar", "Güncellenecek"],
            ["Tablo 2.1.", "Hareket yakalama formülasyonu ve kalite metrikleri", "Güncellenecek"],
            ["Tablo 2.2.", "Donanım bileşenleri ve görevleri", "Güncellenecek"],
            ["Tablo 2.3.", "Yazılım katmanları ve sorumlulukları", "Güncellenecek"],
            ["Tablo 2.4.", "Backend veri modeli ve kalıcılık alanları", "Güncellenecek"],
            ["Tablo 2.5.", "Mobil ekranlar ve kullanıcı akışındaki görevleri", "Güncellenecek"],
            ["Tablo 2.6.", "API uçları ve sistem sözleşmeleri", "Güncellenecek"],
            ["Tablo 2.7.", "Worker işlem hattı aşamaları", "Güncellenecek"],
            ["Tablo 2.8.", "Artifact sözleşmesi ve kullanım amacı", "Güncellenecek"],
            ["Tablo 2.9.", "Veri yaşam döngüsü ve saklama kararları", "Güncellenecek"],
            ["Tablo 3.1.", "Deney ortamı ve çekim senaryoları", "Güncellenecek"],
            ["Tablo 3.2.", "Sanal laboratuvar bileşenleri", "Güncellenecek"],
            ["Tablo 3.3.", "Test aşaması ve kabul kriterleri", "Güncellenecek"],
            ["Tablo 3.4.", "Önerilen deney veri seti", "Güncellenecek"],
            ["Tablo 3.5.", "QA kanıt matrisi", "Güncellenecek"],
            ["Tablo 4.1.", "Proje ve take yönetimi ekran özellikleri", "Güncellenecek"],
            ["Tablo 4.2.", "Çekim modları ve kamera konfigürasyonu", "Güncellenecek"],
            ["Tablo 4.3.", "Export sonuç ekranındaki artifact listesi", "Güncellenecek"],
            ["Tablo 4.4.", "Multi-view kurulum ve diagnostik ekran bilgileri", "Güncellenecek"],
            ["Tablo 4.5.", "Mobil uygulama yazılım katmanları", "Güncellenecek"],
            ["Tablo 5.1.", "Tehdit, risk ve önlem matrisi", "Güncellenecek"],
            ["Tablo 5.2.", "Yükleme ve işleme hattı bütünlük kontrolleri", "Güncellenecek"],
            ["Tablo 5.3.", "Kişisel veri ve gizlilik kontrol listesi", "Güncellenecek"],
            ["Tablo 5.4.", "Üretim güvenliği için önerilen ek kontroller", "Güncellenecek"],
            ["Tablo 6.1.", "Uçtan uca test aşaması ve kabul kriterleri", "Güncellenecek"],
            ["Tablo 6.2.", "Test senaryoları ve kalite sonuçları", "Güncellenecek"],
            ["Tablo 6.3.", "Mimari kazanımlar ve sistem çıktıları", "Güncellenecek"],
            ["Tablo 7.1.", "Proje iş-zaman çizelgesi", "Güncellenecek"],
            ["Tablo 7.2.", "Proje maliyet analizi", "Güncellenecek"],
            ["Tablo 8.1.", "Değişiklik türleri ve kontrol adımları", "Güncellenecek"],
            ["Tablo 8.2.", "Mühendislik standartları uyum matrisi", "Güncellenecek"],
            ["Tablo 8.3.", "Değişiklik yönetimi sorumluluk matrisi", "Güncellenecek"],
            ["Tablo 9.1.", "Hedef pazar ve kullanım senaryoları", "Güncellenecek"],
            ["Tablo 9.2.", "Gelir modeli ve maliyet kalemleri", "Güncellenecek"],
            ["Tablo 9.3.", "Pazara çıkış ve ürünleştirme planı", "Güncellenecek"],
            ["Tablo 9.4.", "Ticarileşme riskleri ve azaltma stratejileri", "Güncellenecek"],
            ["Tablo 9.5.", "Lean Canvas özeti", "Güncellenecek"],
            ["Tablo 9.6.", "SWOT analizi", "Güncellenecek"],
            ["Tablo 9.7.", "12 aylık ticarileşme yol haritası", "Güncellenecek"],
            ["Tablo A.1.", "Rapora eklenmesi önerilen kanıt ve çıktı listesi", "Güncellenecek"],
            ["Tablo A.2.", "Değişiklik kayıt şablonu", "Güncellenecek"],
            ["Tablo A.3.", "Ticarileşme maliyet varsayımı şablonu", "Güncellenecek"],
        ],
        [1.1, 4.8, 1.2],
    )
    document.add_page_break()

    add_heading(document, "ÖZET", front_h_style)
    for text in [
        "Bu bitirme çalışmasında, mobil cihazlar ile video tabanlı işaretleyicisiz hareket yakalama yapabilen Mocap adlı sistem geliştirilmiştir. Sistem, geleneksel hareket yakalama çözümlerinde kullanılan özel kıyafet, işaretleyici, optik stüdyo ve pahalı donanım gereksinimlerini azaltmayı; kullanıcıların iOS veya Android bir cihazla video kaydedip bu videoyu backend ortamında hareket verisine dönüştürebilmesini amaçlamaktadır.",
        "Mocap mimarisi mobil uygulama, yerel çekim deposu, backend API, PostgreSQL tabanlı kalıcı kayıtlar, S3 uyumlu nesne depolama, worker iş kuyruğu ve WHAM/SMPL/SMPLify tabanlı hareket çözüm hattından oluşmaktadır. Mobil uygulama çekim ekranı, yerel önizleme, kayıt durumu, review akışı, upload ekranı, processing status ekranı ve export result ekranlarıyla kullanıcının tüm süreci takip etmesini sağlar. Backend tarafı proje, take, capture session, upload session, processing job ve export file kayıtlarını yönetir. Worker ise kaynak videoyu indirir, FFmpeg ile normalize eder, WHAM/SMPL çözümünü çalıştırır, hareket temizleme ve BVH üretimi yapar, kalite raporu ve pipeline raporu oluşturur.",
        "Çalışmada tek kamera üretim hattı ana güvenilir yol olarak ele alınmış, dual/pro çoklu kamera altyapısı ise eş zamanlama, kalibrasyon, triangulation ve tanısal kalite metrikleri için tasarlanmıştır. Sistem; quality_report, preview_summary, solved_motion, smpl_parameters, cleanup_report, motion_pipeline_report ve BVH gibi artifact dosyaları üreterek hem kullanıcı arayüzünde sonuç gösterimi hem de Blender, Unity, Unreal ve diğer DCC araçlarına aktarım için veri sağlar. Proje sonucunda mobil cihaz odaklı, backend destekli, genişletilebilir ve test edilebilir bir hareket yakalama altyapısı ortaya konmuştur.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    p = document.add_paragraph(style=body_style)
    p.add_run("Anahtar kelimeler: ").bold = True
    p.add_run("Mocap, işaretleyicisiz hareket yakalama, mobil video, WHAM, SMPL, BVH, bulut işleme, çoklu kamera, kalite raporu")
    document.add_page_break()

    add_heading(document, "GİRİŞ", h1_style)
    add_paragraph(
        document,
        "Hareket yakalama, bir insanın vücut hareketlerinin sayısal verilere dönüştürülerek animasyon, oyun, sanal üretim, spor analizi veya araştırma süreçlerinde kullanılmasını sağlayan önemli bir teknolojidir. Geleneksel hareket yakalama sistemleri çoğu zaman özel stüdyo, çok sayıda optik kamera, reflektif işaretleyici, kalibrasyon ekipmanı ve yüksek kurulum maliyeti gerektirir. Bu durum bağımsız geliştiriciler, küçük stüdyolar, öğrenciler ve hızlı prototipleme yapan ekipler için hareket yakalamayı erişilmesi zor bir teknolojiye dönüştürmektedir.",
        first_style,
        first=True,
    )
    for text in [
        "Mocap projesi bu problemi mobil cihazlar üzerinden çözmeyi hedeflemektedir. Kullanıcı, özel bir kıyafet giymeden veya vücuduna işaretleyici yerleştirmeden, mobil cihaz kamerasıyla bir hareket videosu kaydeder. Kaydedilen orijinal video ve çekim metadata bilgileri backend sistemine yüklenir. Backend tarafında video normalize edilir, WHAM/SMPL/SMPLify tabanlı hareket çözüm hattından geçirilir ve sonuç olarak animasyon dosyaları, kalite raporları ve önizleme verileri oluşturulur.",
        "Bu çalışma yalnızca bir kamera ekranı veya basit bir video kaydedici uygulaması değildir. Proje; mobil uygulama, native kamera kayıt modülü, upload yöneticisi, backend API, veritabanı, nesne depolama, iş kuyruğu, worker işlem hattı, hareket temizleme, export doğrulama ve kullanıcıya sunulan sonuç ekranlarını kapsayan bütüncül bir sistemdir. Bu nedenle rapor boyunca Mocap, uçtan uca hareket yakalama platformu olarak ele alınmıştır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "Mobil cihaz -> Video + metadata -> Imzali upload URL -> S3 uyumlu depolama",
            "Backend API -> Processing job -> Worker -> WHAM/SMPL solve -> Cleanup -> BVH + raporlar",
            "Mobil sonuc ekrani <- Export listesi + kalite raporu + indirme URL'leri",
        ],
    )
    add_caption(document, "Şekil 1.1. Mocap genel iş akışı", caption_style)
    add_intro_expansion(document, body_style, caption_style)

    add_heading(document, "Problemin Tanımı: Hareket Yakalamada Erişilebilirlik Sorunu", h2_style)
    for text in [
        "Bu bölümde projenin problem tanımı, temel amacı ve hedef kullanıcı kitlesi açıklanmaktadır. Mocap sisteminin çıkış noktası, hareket yakalama işleminin yaygın kullanıcılar için pahalı, karmaşık ve donanım bağımlı olmasıdır. Animasyon üretmek isteyen bir geliştirici veya öğrenci, klasik optik sistemlerde kamera dizilimi, marker seti, stüdyo aydınlatması ve kalibrasyon sürecini yönetmek zorundadır. Mobil cihazların kamera kalitesi ve video işleme kapasitesi arttıkça bu sürecin daha erişilebilir bir yazılım sistemi ile yönetilmesi mümkün hale gelmiştir.",
        "Projenin ana amacı, mobil cihazla kaydedilen bir insan hareketi videosunu backend ortamında işleyerek animasyon üretiminde kullanılabilir hareket verisine dönüştürmektir. Bu amaç doğrultusunda kullanıcı deneyiminin basit tutulması, çekim verisinin kaybolmaması, video ve metadata bilgisinin birlikte taşınması, işlem durumunun takip edilebilir olması ve oluşan export dosyalarının kalite raporlarıyla birlikte sunulması hedeflenmiştir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 1.1. Problem, amaç ve kapsam matrisi", caption_style)
    add_table(
        document,
        ["Başlık", "Açıklama", "Mocap karşılığı"],
        [
            ["Problem", "Geleneksel hareket yakalama pahalı ve karmaşıktır.", "Mobil video ve backend solve hattı ile erişilebilir yakalama süreci."],
            ["Amaç", "İşaretleyicisiz insan hareketini animasyon verisine dönüştürmek.", "Video + metadata -> WHAM/SMPL -> BVH/JSON artifact üretimi."],
            ["Kullanıcı", "Animatör, oyun geliştirici, öğrenci, küçük stüdyo.", "Capture, review, upload, processing ve export ekranları."],
            ["Kapsam", "Mobil çekim, güvenli upload, bulut işleme, kalite raporu.", "React Native mobil uygulama, Fastify backend, worker, PostgreSQL, S3."],
            ["Sınır", "Tam üretim kalitesinde çoklu kamera solve henüz nihai yol değildir.", "Dual/pro kamera tanısal reconstruction ve gelecek geliştirme altyapısı."],
        ],
        [1.2, 2.5, 2.7],
    )

    add_heading(document, "Projenin Amacı ve Kapsamı", h2_style)
    for text in [
        "Mocap projesinin katkısı, mobil cihazdaki video kaydını backend tabanlı bir hareket yakalama hattıyla birleştirmesidir. Birçok prototip uygulama yalnızca cihaz üzerinde anlık iskelet çizimi veya basit landmark gösterimi yaparken, bu projede üretim kaynağı olarak orijinal video korunur. Orijinal videonun korunması, daha sonra farklı model sürümleriyle yeniden işleme, kalite karşılaştırması yapma ve backend tarafında daha güçlü GPU kaynaklarından yararlanma imkanı sağlar.",
        "Sistemde mobil uygulama kullanıcıya çekim ve takip deneyimi sunarken, asıl ağır işlem backend worker tarafında yürütülür. Bu tercih mobil cihazın batarya, ısınma ve işlemci sınırlarını azaltır. Aynı zamanda WHAM/SMPL gibi lisanslı model asset'leri ve GPU gerektiren bileşenlerin kullanıcı cihazına dağıtılmasını engeller. Böylece model yönetimi, güvenlik ve ölçeklenebilirlik backend ortamında kontrol edilebilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Fonksiyonel kapsam ve kullanıcı akışı", h3_style)
    for text in [
        "Proje kapsamı mobil uygulama, backend API, veritabanı, nesne depolama ve worker işlem hattı olmak üzere beş ana alanda değerlendirilebilir. Mobil uygulama çekim, önizleme, review, upload ve export sonucu görüntüleme görevlerini üstlenir. Backend API proje/take/session/upload/job/export kayıtlarını yönetir. Veritabanı bu kayıtları kalıcı hale getirir. Nesne depolama video ve artifact dosyalarını saklar. Worker ise video normalizasyonu, WHAM/SMPL çözümü, temizleme, doğrulama ve export üretimini gerçekleştirir.",
        "Bu kapsam içinde tek kamera akışı ana üretim hattı olarak kabul edilmiştir. Dual kamera ve pro dört kamera tarafında ise session oluşturma, cihaz rolü, device index, join token, zaman senkronizasyonu, frame eşleştirme, kalibrasyon ve triangulation gibi altyapılar bulunmaktadır. Bu altyapılar final animasyon kalitesini otomatik olarak iyileştiren tamamlanmış bir çoklu kamera solver olarak değil; kalite tanısı ve sonraki fazlarda gerçek çoklu kamera kısıtlarının WHAM/SMPL çözümüne bağlanması için temel olarak konumlandırılmıştır.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Teknik kapsam ve sistem bileşenleri", h3_style)
    for text in [
        "Projenin temel gereksinimleri fonksiyonel ve fonksiyonel olmayan gereksinimler olarak ayrılmıştır. Fonksiyonel gereksinimler; kamera izni alma, native kamera preview gösterme, video kaydetme, çekim metadata'sı üretme, yerel take kaydı oluşturma, backend'e upload yapma, processing job başlatma, job durumunu takip etme ve sonuç export dosyalarını listeleme adımlarını kapsar.",
        "Fonksiyonel olmayan gereksinimler ise veri bütünlüğü, hata durumunda okunabilir mesaj üretimi, retry mekanizması, video dosyalarının doğrudan API sunucusu üzerinden taşınmaması, imzalı URL kullanımı, iş kuyruğu durumlarının izlenebilir olması, kalite raporlarının saklanması ve model asset'lerinin repository dışında tutulmasıdır. Bu gereksinimler projenin yalnızca çalışır bir demo değil, genişletilebilir bir mühendislik sistemi olmasını sağlar.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Varsayımlar, sınırlılıklar ve lisans koşulları", h3_style)
    for text in [
        "Mocap sisteminin varsayımları ve sınırlılıkları açık biçimde belirtilmelidir. Tek kamera WHAM hattı, sistemin ana güvenilir üretim yoludur. Final BVH ve ana solved motion çıktıları mevcut durumda primary camera WHAM solve üzerinden üretilmektedir. Çoklu kamera tarafında bulunan reconstruction artifact'leri, kalite metriği ve diagnostic raporlar final animasyonun kaynağı olarak sunulmamalıdır; bunlar gelecekte daha doğru kamera kısıtları ve gerçek çoklu kamera optimizasyonu için altyapı sağlar.",
        "Ayrıca WHAM repository'si, checkpoint dosyaları ve SMPL gibi lisanslı model asset'leri kaynak kod deposuna dahil edilmemelidir. Rapor içinde de özel environment değerleri, veritabanı bağlantıları, S3 erişim anahtarları, imzalı URL'ler veya özel model dosya yolları paylaşılmamalıdır. Bu sınırlılık, veri güvenliği ve lisans uyumu açısından özellikle önemlidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Beklenen çıktılar ve artifact yapısı", h3_style)
    for text in [
        "Projenin beklenen çıktıları iki grupta toplanır: kullanıcıya dönük çıktılar ve teknik doğrulama çıktıları. Kullanıcıya dönük çıktılar arasında backend tarafından üretilen BVH dosyası, solved motion JSON, WHAM overlay preview videosu, kalite özeti ve indirme bağlantıları yer alır. Teknik doğrulama çıktıları arasında quality_report, preview_summary, cleanup_report, motion_pipeline_report, SMPL parameters ve job timeline event kayıtları bulunur.",
        "Rapor hazırlanırken bu çıktılar yalnızca isim olarak yazılmamalı; her çıktının neyi doğruladığı, hangi durumda başarılı kabul edildiği ve hangi kullanıcı ihtiyacına karşılık geldiği açıklanmalıdır. Örneğin BVH dosyası Blender'a import edilebiliyorsa bu, export formatının yapısal olarak doğru olduğuna dair kanıt sağlar. Quality report içindeki skor ve uyarılar ise hareketin sanatsal olarak kusursuz olduğunu değil, sistemin ölçebildiği teknik kalite göstergelerini ifade eder.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Önerilen Sistemin Ana Hatları", h2_style)
    add_paragraph(
        document,
        "Rapor dokuz ana bölüm halinde düzenlenmiştir. Birinci bölüm (GİRİŞ) projenin problemini, amacını ve kapsamını tanıtır. İkinci bölüm (SİSTEMATİK YAKLAŞIM) hareket yakalama formülasyonunu ve sistemin mobil, backend, worker ve bulut katmanlarını ayrıntılandırır. Üçüncü bölüm (DENEY DÜZENEĞİ VE SANAL LABORATUVAR) kullanım senaryolarını ve test ortamını açıklar. Dördüncü bölüm (KULLANICI ARAYÜZÜ) mobil uygulamanın tüm ekran akışını ve backend entegrasyonunu belgeler. Beşinci bölüm (VERİ GÜVENLİĞİ DEĞERLENDİRMESİ) veri bütünlüğü, güvenli yükleme, model varlık güvenliği ve kişisel veri risklerini ele alır. Altıncı bölüm (TEST VE SONUÇLAR) test yöntemi, performans/kalite sonuçları ve teknik bulgular sunar. Yedinci bölüm (PROJE PLANI VE MALİYET ANALİZİ) iş-zaman çizelgesi ve maliyet kalemlerini inceler. Sekizinci bölüm (DEĞİŞİKLİK YÖNETİMİ VE MÜHENDİSLİK STANDARTLARI) yazılım geliştirme sürecini ve kalite kapılarını açıklar. Dokuzuncu ve son bölüm (TİCARİLEŞME PLANI) pazar analizi, gelir modeli ve büyüme stratejisini sunar.",
        body_style,
        first=True,
    )
    document.add_page_break()

    add_heading(document, "SİSTEMATİK YAKLAŞIM", h1_style)
    add_paragraph(
        document,
        "Mocap sisteminin sistematik yaklaşımı, hareket yakalama sürecini bağımsız fakat birbiriyle sözleşmeli bileşenlere ayırmaya dayanır. Mobil uygulama görüntüyü kaydeder ve metadata üretir; backend API iş kayıtlarını yönetir; nesne depolama büyük dosyaları taşır; worker ağır model çıkarımını yapar; sonuç ekranları export dosyalarını ve kalite raporlarını kullanıcıya sunar. Bu bölümde sistemin matematiksel formülasyonu, donanım mimarisi, yazılım mimarisi ve bulut mimarisi açıklanmaktadır.",
        first_style,
        first=True,
    )
    add_diagram_box(
        document,
        [
            "Capture UI | Native Camera | Local Take Repo",
            "        v",
            "Signed Upload Manager -> Backend API -> PostgreSQL",
            "        v                         v",
            "S3/MinIO object storage <- Worker queue -> WHAM/SMPL worker",
            "        v",
            "Export Result UI: BVH, JSON raporlar, overlay preview",
        ],
    )
    add_caption(document, "Şekil 2.1. Sistem mimarisi blok diyagramı", caption_style)

    add_heading(document, "Video Tabanlı Hareket Yakalama Formülasyonu", h2_style)
    for text in [
        "Mocap projesinde temel problem, bir veya birden fazla mobil cihaz tarafından kaydedilen video dizilerinden animasyon üretiminde kullanılabilir, zaman içinde tutarlı ve doğrulanabilir insan hareketi verisi elde etmektir. Tek kamera senaryosunda giriş V(t) ile gösterilen bir video dizisi ve M ile gösterilen capture metadata paketidir. Çıkış ise A ile gösterilen animasyon artifact setidir. Bu artifact seti BVH, solved_motion_json, smpl_parameters_json, quality_report_json ve motion_pipeline_report_json gibi dosyalardan oluşur.",
        "Çoklu kamera senaryosunda giriş, cihaz indeksine göre V0(t), V1(t), V2(t), V3(t) video dizileri ve her cihaz için M0, M1, M2, M3 metadata paketleridir. Bu durumda sistem önce zaman eşleştirmesi ve kamera kalibrasyon verilerini kullanarak tanısal multi-view reconstruction çıktıları üretir. Ancak mevcut üretim yaklaşımında final animasyon kaynağı çoğunlukla primary WHAM solve olarak kalır. Bu karar, yanlış kalibrasyon veya eksik senkronizasyon durumunda sahte başarı üretilmesini engeller.",
        "Kalite değerlendirmesi tek bir sayıya indirgenmemelidir. Sistemde frame count, duration, FPS, root motion stabilitesi, NaN/Infinity bulunmaması, eksik eklem rotasyonu, BVH hiyerarşi doğruluğu, Blender import sonucu, multi-view eşleşme sayısı, sync confidence, reprojection error ve triangulated landmark ratio gibi metrikler birlikte değerlendirilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 2.1. Hareket yakalama formülasyonu ve kalite metrikleri", caption_style)
    add_table(
        document,
        ["Terim/Metrik", "Tanım", "Rapora eklenmesi gereken değer"],
        [
            ["V(t)", "Kaynak video kareleri dizisi.", "Video süresi, çözünürlük, FPS, codec, dosya boyutu."],
            ["M", "Capture metadata paketi.", "schema, takeId, deviceId, captureMode, timestamp, camera ve quality alanları."],
            ["A", "Üretilen artifact seti.", "BVH, solved motion, SMPL parameters, quality report, preview summary."],
            ["Q", "Teknik kalite skoru.", "quality_report içindeki score, grade, warnings ve errors."],
            ["Δt", "Çoklu kamera frame eşleşme zaman farkı.", "averageTimeDeltaMs, p95TimeDeltaMs, syncConfidence."],
            ["E reproj", "Triangulation sonrası yeniden izdüşüm hatası.", "reprojectionErrorPx ve reprojectionP95Px."],
            ["C tri", "Triangulated landmark coverage.", "triangulatedLandmarkRatio ve fallbackLandmarkRatio."],
        ],
        [1.2, 2.5, 2.8],
    )

    add_heading(document, "Genel Sistem Mimarisi", h2_style)
    for text in [
        "Mocap sisteminin donanım mimarisi kullanıcı cihazı, ağ bağlantısı, backend sunucusu, nesne depolama ve worker çalışma ortamı bileşenlerinden oluşur. Kullanıcı tarafında iOS veya Android cihaz kamerası video kaynağıdır. Mobil cihazda native kamera modülü, AVFoundation veya CameraX tabanlı kayıt yapar ve orijinal video dosyasını yerel depoda tutar. Çekim sırasında cihazın kamera pozisyonu, video yönü, FPS ve dosya boyutu metadata paketine işlenir.",
        "Backend tarafında CPU ağırlıklı API sunucusu ve GPU gerektirebilen worker ortamı ayrıştırılmıştır. API sunucusu büyük video dosyasını doğrudan almak yerine imzalı upload URL üretir. Bu sayede video nesne depolamaya doğrudan gider. Worker ise gerekli olduğunda bu videoyu indirir, normalize eder ve WHAM/SMPL model hattını çalıştırır. Bu ayrım, mobil cihazdaki yükü azaltırken backend tarafında ölçeklenebilirlik sağlar.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 2.2. Donanım bileşenleri ve görevleri", caption_style)
    add_table(
        document,
        ["Bileşen", "Görev", "Raporda gösterilecek kanıt"],
        [
            ["Mobil cihaz", "Video kaydı, kamera preview, kullanıcı etkileşimi.", "Capture ekran görüntüsü, video metadata örneği."],
            ["Ağ bağlantısı", "Upload ve WebSocket/HTTP haberleşmesi.", "Upload progress ekranı, job polling örneği."],
            ["Backend sunucusu", "API, session, upload, processing ve export yönetimi.", "Endpoint listesi ve servis katmanı tablosu."],
            ["PostgreSQL", "Project, take, capture session, job ve export kayıtları.", "Veri modeli tablosu, migration özeti."],
            ["S3 uyumlu depolama", "Video, metadata, normalized video ve artifact dosyaları.", "Object key şeması ve güvenlik akışı."],
            ["GPU worker", "WHAM/SMPL solve, cleanup, BVH ve rapor üretimi.", "Pipeline report ve worker stage akışı."],
        ],
        [1.4, 2.5, 2.6],
    )

    add_heading(document, "Mobil Uygulama ve Donanım Katmanı", h2_style)
    for text in [
        "Yazılım mimarisi katmanlı bir yapıdadır. Mobil uygulamada app shell, feature ekranları, feature hook'ları, domain modelleri ve infrastructure adaptörleri ayrılmıştır. CaptureScreen ve MultiViewSetupScreen kullanıcı etkileşimini yönetirken, useWhamCapture ve useRecorder native kamera kaydı ve capture metadata üretimini koordine eder. Local take repository, meta.json ve JSONL chunk dosyalarıyla cihaz üzerindeki kayıtları saklar.",
        "Backend tarafında Fastify route'ları ince tutulmuş, asıl iş mantığı ProjectService, TakeService, CaptureSessionService, UploadService, ProcessingService ve ExportService içinde toplanmıştır. Repository katmanı PostgreSQL sorgularını kapsar. Worker tarafı ise processing job'ı alır, video kaynaklarını hazırlar, çoklu kamera tanısal aşamasını gerektiğinde çalıştırır, WHAM/SMPL solve yapar, hareket temizleme ve export doğrulama adımlarını yürütür.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "Screens -> Hooks/State -> Domain models -> Infra adapters",
            "Routes  -> Services    -> Repositories -> PostgreSQL",
            "Worker  -> Video pipeline -> WHAM/SMPL -> Export validation",
        ],
    )
    add_caption(document, "Şekil 2.2. Mobil, backend ve worker veri akışı", caption_style)
    add_caption(document, "Tablo 2.3. Yazılım katmanları ve sorumlulukları", caption_style)
    add_table(
        document,
        ["Katman", "Temel dosya/sorumluluk", "Açıklama"],
        [
            ["Mobil ekranlar", "Capture, Review, Projects, Upload, Processing, ExportResult", "Kullanıcının çekimden sonuç izlemeye kadar tüm akışını yönetir."],
            ["State/hook", "captureStore, multiViewStore, useRecorder, useWhamCapture", "Canlı durum, kayıt durumu, çoklu kamera bağlantısı ve metadata üretimini koordine eder."],
            ["Domain", "Take, CaptureMetadata, PoseFrame, export ve cleanup pipeline", "Projenin temel veri modelleri ve hareket işleme kavramları burada tanımlıdır."],
            ["Infrastructure", "MocapApiClient, SignedUrlUploadManager, TakeRepo.fs", "Backend, yerel dosya sistemi ve upload altyapısına erişim sağlar."],
            ["Backend API", "Fastify routes ve servisler", "Proje, take, session, upload, processing job ve export kayıtlarını yönetir."],
            ["Worker", "processJob, premiumMotionSolver, motionCleanup, bvhWriter", "Videoyu işleyerek hareket ve rapor artifact'lerini üretir."],
        ],
        [1.5, 2.5, 2.7],
    )
    add_heading(document, "Backend API, Veri Modeli ve PostgreSQL Yapısı", h2_style)
    add_caption(document, "Tablo 2.4. Backend veri modeli ve kalıcılık alanları", caption_style)
    add_table(
        document,
        ["Tablo/Kayıt", "Sorumluluk", "Önemli alanlar"],
        [
            ["projects", "Kullanıcıya ait proje grupları.", "id, user_id, name, created_at."],
            ["takes", "Backend çekim kaydı ve işlem durumu.", "capture_mode, expected_video_count, status."],
            ["capture_sessions", "Dual/pro çoklu cihaz oturumları.", "join_token, expected_device_count, status, sync_metadata."],
            ["capture_videos", "Her cihaz için upload edilmiş video ve metadata.", "device_index, device_role, video_storage_key, capture_metadata."],
            ["processing_jobs", "Worker kuyruğu ve job state makinesi.", "state, preset, progress, error_code."],
            ["export_files", "Üretilen dosya kayıtları.", "format, artifact_name, storage_key, file_size_bytes."],
        ],
        [1.4, 2.6, 2.7],
    )

    add_heading(document, "Worker Pipeline ve Model Çalışma Ortamı (WHAM/SMPL)", h2_style)
    add_paragraph(
        document,
        "Worker pipeline Mocap sisteminin ağır hesaplama katmanıdır. Bir processing job kabul edildiğinde worker önce kaynak video ve metadata dosyalarını nesne depolamadan indirir. FFmpeg ile normalize aşamasında video sabit FPS ve çözünürlüğe getirilir; bu adım WHAM modelinin tutarlı giriş almasını sağlar. Ardından WHAM/SMPL solve aşamasında her kare için pose parametreleri ve root motion tahmin edilir. Hareket temizleme aşamasında jitter giderme, velocity sınırlama ve NaN kontrolü yapılır. Son olarak BVH yazımı ve kalite raporu üretimi gerçekleştirilir.",
        body_style,
        first=True,
    )
    add_paragraph(
        document,
        "WHAM modeli video dizisinden 3B insan vücut hareketi ve SMPL parametrelerini çıkarır. Bu model araştırma lisansıyla kullanılmakta olup runtime için uygun Python ortamı, checkpoint dosyaları ve SMPL model asset'leri gerekmektedir. Worker pipeline bu varlıkları kaynak kod deposuna dahil etmeden, güvenli bir çalışma ortamında yönetir. Ticari kullanım için WHAM ve SMPL ayrı lisans anlaşması gerektirir; bu durum raporun veri güvenliği ve lisans bölümünde açıklanmaktadır.",
        body_style,
        first=True,
    )
    add_heading(document, "Bulut, S3 Nesne Depolama ve İş Kuyruğu Mimarisi", h2_style)
    for text in [
        "Mocap'ın bulut mimarisi, büyük video dosyalarının güvenli ve ölçeklenebilir biçimde işlenebilmesi için tasarlanmıştır. Mobil cihaz, video dosyasını backend API'ye doğrudan göndermek yerine backend'den imzalı upload URL alır ve dosyayı S3 uyumlu nesne depolamaya yükler. Backend yalnızca upload session, object key, metadata ve dosya boyutu doğrulamalarını yönetir. Bu yaklaşım API sunucusunun bant genişliği yükünü azaltır.",
        "Processing job oluşturulduğunda worker, PostgreSQL üzerinden kuyruğa alınmış işi claim eder veya RunPod gibi serverless GPU ortamına dispatch edilen job'ı işler. Worker kaynak videoyu indirir, FFmpeg ile normalize eder, WHAM/SMPL çözümünü çalıştırır ve oluşan artifact dosyalarını tekrar nesne depolamaya yazar. Sonuçlar export_files tablosuna kaydedilir ve mobil uygulama yalnızca geçici signed download URL ile bu dosyalara erişir.",
        "Bulut mimarisinde üretim ortamı ile yerel geliştirme ortamı birbirinden ayrılmıştır. Yerel geliştirmede Docker Compose ile PostgreSQL ve MinIO kullanılabilir. Gerçek WHAM/SMPL işlemleri için model repository'si, checkpoint dosyaları, lisanslı SMPL asset'leri ve uygun Python/GPU çalışma ortamı gereklidir. Bu varlıklar kaynak kod deposunda tutulmamalıdır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "POST /uploads/init -> signed PUT urls",
            "Mobile PUT metadata.json + source video -> S3 compatible storage",
            "POST /uploads/complete -> size + metadata validation",
            "POST /process -> queued job -> worker -> export files",
        ],
    )
    add_caption(document, "Şekil 2.3. Bulut işleme ve artifact üretim hattı", caption_style)
    add_system_expansion(document, body_style, caption_style)
    document.add_page_break()

    add_heading(document, "DENEY DÜZENEĞİ VE SANAL LABORATUVAR", h1_style)
    add_paragraph(
        document,
        "Bu bölümde Mocap sisteminin uçtan uca kullanım senaryoları, fiziksel deney düzeneği ve yerel/sanal laboratuvar ortamı açıklanmaktadır. Projenin doğru değerlendirilmesi için yalnızca uygulamanın açılması yeterli değildir. Kamera preview, video kaydı, metadata üretimi, upload, backend job state geçişleri, worker artifact üretimi, kalite raporu ve BVH import doğrulaması birlikte test edilmelidir.",
        first_style,
        first=True,
    )

    add_heading(document, "Uygulama İş Akışı (User Flow) ve Kullanım Senaryoları", h2_style)
    for text in [
        "Mocap sisteminin temel kullanım senaryosu şu adımlardan oluşur: Kullanıcı mobil uygulamada proje oluşturur, çekim modunu seçer (solo, dual veya pro), kamera kaydına başlar. Çekim tamamlandığında Review ekranında video önizlemesi ve metadata özeti görülür. Kullanıcı upload düğmesine bastığında mobil uygulama signed URL üzerinden metadata ve video dosyasını backend'e iletir. Backend upload doğrulamasının ardından processing job oluşturur. İşlem tamamlandığında kullanıcı export sonuç ekranında BVH dosyasını, kalite raporunu ve önizleme videosunu inceler.",
        "Çoklu kamera senaryosunda host cihaz oturumu başlatır ve join token üretir. Guest cihazlar bu token'ı girerek oturuma bağlanır; her cihazın rolü ve indeksi backend tarafından atanır. Tüm beklenen cihazlar bağlandığında kayıt eş zamanlı başlatılır. İşlem sonunda multi-view diagnostics bölümünde matched frame count, sync confidence ve reconstruction availability gösterilir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "Proje olustur -> Cekim modu sec -> Kamera kaydi baslat",
            "Review -> Upload (signed URL) -> Processing job",
            "Processing status ekrani -> Export sonuc ekrani",
            "BVH indir | Kalite raporu | Preview video",
        ],
    )
    add_caption(document, "Şekil 3.1. Temel kullanım senaryosu iş akışı", caption_style)

    add_heading(document, "Tek ve Çoklu Kamera Deney Düzeneği", h2_style)
    for text in [
        "Tek kamera deneyinde bir mobil cihaz sabit veya elde tutulur konumda oyuncuyu görecek şekilde yerleştirilir. Oyuncunun tüm vücudunun kadrajda kalması, ortam ışığının yeterli olması, arka planın fazla karmaşık olmaması ve çekim boyunca kamera titremesinin sınırlı tutulması beklenir. Çekimden sonra video metadata'sında çözünürlük, FPS, dosya boyutu, süre, kamera yönü ve timestamp bilgileri yer almalıdır.",
        "Dual kamera deneyinde iki cihaz aynı hareketi farklı açılardan kaydeder. Host ve guest cihazlar backend capture session üzerinden eşleşir; join token, device role ve device index bilgileri kayıt altına alınır. Bu modda amaç yalnızca iki videonun varlığını göstermek değil; zaman senkronizasyonu, camera role, frame eşleştirme ve reconstruction diagnostic metriklerinin üretilebildiğini göstermektir.",
        "Pro dört kamera deneyinde front, right, back ve left rolleriyle dört cihazın aynı capture session'a bağlanması hedeflenir. Bu mod raporda tamamlanmış final motion iyileştirme iddiası olarak sunulmamalı; çoklu kamera altyapısı, cihaz kayıt semantiği, beklenen video sayısı ve ileride gerçek çoklu kamera solve için hazırlanmış platform olarak anlatılmalıdır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "Tek kamera:  [Telefon] ----> [Oyuncu]",
            "Dual kamera: [Telefon A] -> [Oyuncu] <- [Telefon B]",
            "Pro 4:       Front / Right / Back / Left kamera rolleri",
        ],
    )
    add_caption(document, "Şekil 3.2. Tek kamera ve çoklu kamera deney düzeneği", caption_style)
    add_caption(document, "Tablo 3.1. Deney ortamı ve çekim senaryoları", caption_style)
    add_table(
        document,
        ["Senaryo", "Kurulum", "Ölçülecek/eklenecek kanıt"],
        [
            ["Solo capture", "Bir mobil cihaz, tüm vücut görünür, 10-20 sn hareket.", "Capture ekranı, video metadata, upload progress, job success, BVH çıktısı."],
            ["Dual capture", "İki cihaz, host/guest rolü, aynı LAN/backend relay.", "Join token, device index, sync confidence, matched frame count, multi-view diagnostics."],
            ["Pro 4 camera", "Dört cihaz, front/right/back/left rolleri.", "Capture session cihaz listesi, expectedDeviceCount=4, upload bekleme davranışı."],
            ["Kalite varyasyonu", "İyi ışık, düşük ışık, hızlı hareket, kısmi görünmezlik.", "Quality report score, warnings, tracking/fallback açıklamaları."],
        ],
        [1.4, 2.5, 2.9],
    )

    add_heading(document, "Sanal Laboratuvar ve Yerel Geliştirme Ortamı", h2_style)
    for text in [
        "Sanal laboratuvar, gerçek mobil cihaz ve gerçek model runtime olmadan da sistemin önemli bölümlerinin test edilebildiği geliştirme ortamıdır. Yerel backend API, PostgreSQL ve MinIO ile çalıştırılabilir. Mobil uygulama yerel veya paylaşımlı backend'e bağlanarak proje oluşturma, take oluşturma, signed upload alma, upload tamamlama ve processing job başlatma adımlarını test edebilir.",
        "Worker tarafında gerçek WHAM runtime gerektiren testler ayrı değerlendirilmelidir. WHAM repository'si, checkpoint dosyaları, SMPL asset'leri ve Python bağımlılıkları yoksa tam üretim solve çalıştırılamaz; buna rağmen metadata validation, upload state, job state, export schema ve bazı synthetic/golden QA testleri çalıştırılabilir. Rapor hazırlanırken hangi testlerin gerçek modelle, hangilerinin fixture veya synthetic veriyle yapıldığı açıkça yazılmalıdır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 3.2. Sanal laboratuvar bileşenleri", caption_style)
    add_table(
        document,
        ["Bileşen", "Amaç", "Rapora eklenmesi gereken çıktı"],
        [
            ["Local backend", "API endpoint ve state akışını denemek.", "GET /health sonucu, endpoint listesi, örnek job timeline."],
            ["PostgreSQL", "Kalıcı kayıt ve worker kuyruğu.", "projects/takes/jobs/export_files kayıt örnekleri."],
            ["MinIO/S3", "Video ve artifact nesne depolama.", "Bucket yapısı, object key örnekleri, upload tamamlandı kanıtı."],
            ["Worker", "Video normalize, WHAM solve, cleanup, export.", "motion_pipeline_report, quality_report, BVH dosyası."],
            ["Mobil uygulama", "Kullanıcı akışını uçtan uca görmek.", "Capture, Upload, Processing, ExportResult ekran görüntüleri."],
        ],
        [1.4, 2.5, 2.9],
    )

    add_heading(document, "Uçtan Uca Test Aşaması ve Kalite Doğrulaması", h2_style)
    for text in [
        "Test aşaması üç seviyede yürütülmelidir: mobil kullanıcı akışı testleri, backend/worker doğrulama testleri ve artifact kalite testleri. Mobil testlerde kamera izni, preview, kayıt başlatma/durdurma, yerel take oluşturma, metadata doğrulama, upload progress ve sonuç ekranı kontrol edilir. Backend testlerinde upload session, expected video count, capture mode, job state ve export list endpoint'leri doğrulanır. Worker testlerinde video normalize etme, WHAM/SMPL solve, cleanup, BVH üretimi ve kalite raporu kontrol edilir.",
        "Rapor içinde testler yalnızca 'test edildi' cümlesiyle bırakılmamalıdır. Her test için giriş verisi, beklenen sonuç, gözlenen sonuç, başarı kriteri ve varsa ekran görüntüsü veya log kanıtı verilmelidir. Özellikle kalite raporu, motion pipeline raporu ve Blender import sonucu test kanıtı olarak değerlidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "queued -> ingesting -> extracting_frames -> solving_motion -> cleaning -> exporting",
            "          |                                                        |",
            "          +-------------------- failed/canceled -------------------+",
            "exporting -> succeeded -> ExportResult screen",
        ],
    )
    add_caption(document, "Şekil 3.3. Processing job durum akışı", caption_style)
    add_caption(document, "Tablo 3.3. Test aşaması ve kabul kriterleri", caption_style)
    add_table(
        document,
        ["Test", "Başarı kriteri", "Rapora eklenecek kanıt"],
        [
            ["Kamera kaydı", "Video dosyası oluşur, metadata geçerlidir.", "Video metadata tablosu, Capture/Preview ekran görüntüsü."],
            ["Upload", "Metadata ve video signed URL ile yüklenir.", "Upload progress ekranı, tamamlanan upload session kaydı."],
            ["Processing", "Job terminal state'e ulaşır.", "Job timeline tablosu, state geçişleri."],
            ["Artifact üretimi", "BVH ve JSON raporlar export_files'a kaydedilir.", "Export listesi ve artifact adları."],
            ["Kalite raporu", "Score/grade/warnings/errors okunabilir.", "quality_report özet tablosu."],
            ["BVH import", "Blender yapısal import hatası vermez.", "Blender ekran görüntüsü veya smoke test sonucu."],
            ["Çoklu kamera diagnostic", "Eksik kalibrasyonda sahte başarı üretilmez.", "missing_calibration veya diagnostic_only uyarıları."],
        ],
        [1.4, 2.5, 2.9],
    )
    add_experiment_expansion(document, body_style, caption_style)
    document.add_page_break()

    add_ui_section(document, h1_style, h2_style, h3_style, first_style, body_style, caption_style)
    document.add_page_break()

    add_heading(document, "VERİ GÜVENLİĞİ DEĞERLENDİRMESİ", h1_style)
    add_paragraph(
        document,
        "Mocap sistemi video, metadata, model çıktısı ve animasyon dosyası gibi farklı hassasiyet seviyelerine sahip veriler üretir. Kullanıcının vücut hareketini içeren video kişisel veri niteliği taşıyabilir; backend metadata'sı cihaz bilgisi, zaman bilgisi ve çekim parametrelerini içerir; model çıktıları ise kullanıcının hareket profilini temsil eder. Bu nedenle sistemde veri bütünlüğü, erişim kontrolü, geçici URL yönetimi, gizli anahtarların korunması ve işlem izlenebilirliği birlikte değerlendirilmelidir.",
        first_style,
        first=True,
    )

    add_heading(document, "Video ve Metadata Bütünlüğü", h2_style)
    for text in [
        "Mobil çekim sırasında batarya düşmesi, cihazın ısınması, uygulamanın arka plana alınması, depolama alanının dolması veya kamera kaydının beklenmeyen şekilde durması video dosyasında bozulma oluşturabilir. Bu tür bozulmalar doğrudan hareket çözüm kalitesini etkiler. Eksik video süresi, hatalı FPS, uyumsuz frame count veya localUri kaybı upload aşamasının başarısız olmasına neden olabilir.",
        "Mocap mimarisinde bu risk, capture metadata validasyonu ve upload öncesi dosya varlığı kontrolü ile azaltılır. Mobil upload yöneticisi video localUri alanını kontrol eder, dosyanın gerçekten var olup olmadığını sorgular ve dosya boyutunu backend'e bildirir. Backend upload complete aşamasında video ve metadata nesnelerinin beklenen boyutta depolamada bulunduğunu kontrol eder. Böylece eksik veya bozulmuş kaynakların sessizce processing aşamasına geçmesi engellenir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 5.1. Tehdit, risk ve önlem matrisi", caption_style)
    add_table(
        document,
        ["Tehdit/Risk", "Olası etki", "Mocap önlemi"],
        [
            ["Eksik video dosyası", "Upload başlatılamaz veya worker kaynak bulamaz.", "Mobilde file exists kontrolü, backend'de object size doğrulama."],
            ["Metadata uyuşmazlığı", "Yanlış capture mode veya device index ile işleme.", "Mobil ve backend validateCaptureMetadata kontrolleri."],
            ["Signed URL sızıntısı", "Geçici yetkisiz dosya erişimi.", "URL'lerin log/rapora yazılmaması, kısa TTL kullanımı."],
            ["Gizli anahtar sızıntısı", "Depolama veya veritabanı erişimi riske girer.", ".env dosyalarının repository dışında tutulması."],
            ["Model asset lisans ihlali", "Hukuki ve güvenlik riski.", "WHAM/SMPL asset'lerinin repo dışında ve kontrollü ortamda tutulması."],
            ["Sahte çoklu kamera başarısı", "Yanlış kalite iddiası.", "missing_calibration, diagnostic_only ve primary fallback reason alanları."],
        ],
        [1.5, 2.3, 3.0],
    )

    add_heading(document, "Güvenli Yükleme ve Signed URL Yönetimi", h2_style)
    for text in [
        "Bu başlık proje bağlamında cihaz ve ağ kararsızlığının haberleşme kalitesi üzerindeki etkisi olarak ele alınmıştır. Mobil cihazların kablosuz ağ bağlantısı değişebilir, upload sırasında bağlantı kopabilir, signed URL süresi dolabilir veya WebSocket bağlantısı kesilebilir. Bu durumlar kullanıcı deneyimini ve çoklu kamera eşleşmesini doğrudan etkiler.",
        "Sistemde upload yöneticisi hata durumunda fresh signed URL ile tekrar deneme yapacak şekilde tasarlanmıştır. Backend API hataları yapılandırılmış error code, message, requestId ve retryable bilgileriyle döner. Processing status ekranı job durumunu periyodik olarak sorgular ve hata/cancel/retry davranışlarını kullanıcıya sunar. Çoklu kamera tarafında connectionState, timeSyncReady, syncRtt, remote frame age ve matched frame count gibi göstergeler haberleşme kalitesini izlemek için kullanılmalıdır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo 5.2. Yükleme ve işleme hattı bütünlük kontrolleri", caption_style)
    add_table(
        document,
        ["Aşama", "Kontrol", "Başarısızlıkta beklenen davranış"],
        [
            ["Upload init", "deviceIndex, deviceRole, captureMode ve expectedVideoCount uyumu.", "400/409 hata; upload session başlatılmaz."],
            ["Metadata upload", "application/json ve beklenen metadata boyutu.", "Retry veya failed remote state."],
            ["Video upload", "Dosya varlığı, contentType, dosya boyutu.", "Retry veya missing_file/video_upload_failed."],
            ["Upload complete", "ObjectStorage assertObject ve metadata route uyumu.", "Processing başlatılmaz."],
            ["Processing create", "Tüm beklenen videolar yüklenmiş olmalı.", "All expected capture videos uyarısı."],
            ["Worker solve", "Solved motion ve BVH validation.", "Job failed, error_code ve metrics kaydı."],
        ],
        [1.4, 2.6, 2.8],
    )

    add_heading(document, "WHAM/SMPL Model Varlık Güvenliği ve API Güvenliği", h2_style)
    for text in [
        "Bu bölümde WHAM/SMPL model varlıklarının güvenliği, projede kullanılan hareket çözüm hattının sistem güvenilirliği ve kalite değerlendirmesi üzerindeki etkisi olarak ele alınmıştır. WHAM, videodan 3B insan hareketi ve SMPL parametreleri üretmek için kullanılan ana model hattıdır. Bu hattın doğru çalışabilmesi için normalize edilmiş video, uygun model runtime, checkpoint dosyaları, SMPL asset'leri ve Python/GPU ortamı gereklidir.",
        "WHAM hattı başarısız olduğunda sistemin yanlış animasyon üretip kullanıcıya başarı göstermesi kabul edilemez. Bu nedenle worker solved motion validation, SMPL parameter varlığı, BVH hiyerarşi kontrolü, NaN/Infinity kontrolü ve kalite raporu üretimi yapar. Rapor içinde WHAM çıktıları anlatılırken modelin araştırma ve lisans koşulları, runtime gereksinimleri ve final animasyonun tek kamera WHAM'a dayandığı açıkça belirtilmelidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Kimlik doğrulama, yetkilendirme ve gizli değer yönetimi", h2_style)
    for text in [
        "Mevcut backend auth yapısı geliştirme ve prototip düzeyindedir; bearer token değeri kullanıcı id gibi yorumlanmakta ve üretim seviyesinde kimlik doğrulama sağlamamaktadır. Bu durum raporda açıkça teknik borç olarak belirtilmelidir. Üretim ortamında gerçek kullanıcı kimliği, token doğrulama, rol tabanlı erişim, CORS kısıtlaması, rate limit, audit log ve veri saklama politikaları eklenmelidir.",
        "Gizli değerlerin korunması da ayrı bir güvenlik faktörüdür. Veritabanı URL'si, S3 access key, secret key, RunPod API key, signed upload/download URL, model asset yolu ve lisanslı checkpoint dosyaları rapora veya kaynak kod deposuna eklenmemelidir. Rapor kanıtlarında yalnızca maskelemiş object key yapısı, endpoint adları ve örnek şema alanları kullanılmalıdır.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Veri aktarım hızı, signed URL ve object storage etkisi", h2_style)
    for text in [
        "Video tabanlı hareket yakalama sistemlerinde veri hızı doğrudan kullanıcı deneyimini belirler. Yüksek çözünürlüklü video dosyaları mobil ağda uzun upload sürelerine neden olabilir. Bu nedenle Mocap, büyük video dosyasını API sunucusu üzerinden geçirmek yerine doğrudan nesne depolamaya imzalı URL ile yükler. Backend yalnızca kontrol ve kayıt yönetimi yapar.",
        "Rapor içinde upload süresi, video boyutu, FPS, çözünürlük, normalize edilmiş video boyutu, worker işlem süresi ve export dosya boyutları ayrı ayrı verilmelidir. Bu veriler sistemin ölçeklenebilirliği hakkında gerçekçi yorum yapmayı sağlar. Örneğin 10 saniyelik bir solo çekim ile 4 kameralı pro çekim aynı ağ ve depolama maliyetine sahip değildir. Bu fark maliyet analizi ve performans değerlendirmesinde gösterilmelidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)

    add_heading(document, "Kişisel Verilerin Korunması ve KVKK Uyumu", h2_style)
    for text in [
        "Mobil bozulma, çekim sırasında uygulamanın kapanması, frame timestamp alanlarının eksik kalması, cihaz saatlerinin uyuşmaması, ağ bağlantısının kopması veya kamera metadata'sının eksik üretilmesi gibi durumları kapsar. Bu eksikler özellikle dual/pro kamera tarafında sync confidence ve reconstruction kalitesini düşürür. Worker, eksik metadata durumunda fake başarı üretmek yerine missing_timestamps, sync_metadata_incomplete, missing_calibration veya diagnostic_only gibi uyarılar üretmelidir.",
        "Raporun veri güvenliği bölümünde mobil veri açığı yalnızca siber saldırı olarak değil, veri kalitesi ve izlenebilirlik açığı olarak da değerlendirilmelidir. Çünkü hareket yakalama sonucunun doğruluğu, kaynak videonun bütünlüğü kadar metadata'nın doğruluğuna da bağlıdır. Bu nedenle capture metadata alanları raporda tablo halinde gösterilmeli; hangi alanların zorunlu, hangilerinin opsiyonel ve hangilerinin çoklu kamera için kritik olduğu belirtilmelidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_diagram_box(
        document,
        [
            "Local video + metadata -> sanitize localUri -> signed upload",
            "Object size validation -> job timeline",
            "Quality report + motion pipeline report",
        ],
    )
    add_caption(document, "Şekil 5.1. Veri güvenliği ve bütünlük kontrol akışı", caption_style)
    add_security_expansion(document, body_style, caption_style)
    document.add_page_break()

    add_results_section(document, h1_style, h2_style, first_style, body_style, caption_style)
    document.add_page_break()

    add_project_plan_section(document, h1_style, h2_style, first_style, body_style, caption_style)
    document.add_page_break()

    add_change_management_section(document, h1_style, h2_style, first_style, body_style, caption_style)
    document.add_page_break()

    add_commercialization_section(document, h1_style, h2_style, first_style, body_style, caption_style)
    document.add_page_break()

    add_heading(document, "KAYNAKLAR", front_h_style)
    references = [
        "[1] Shin, S., Kim, J., Halilaj, E., Black, M. J., WHAM: Reconstructing World-grounded Humans with Accurate 3D Motion, CVPR 2024. https://openaccess.thecvf.com/content/CVPR2024/papers/Shin_WHAM_Reconstructing_World-grounded_Humans_with_Accurate_3D_Motion_CVPR_2024_paper.pdf",
        "[2] Loper, M., Mahmood, N., Romero, J., Pons-Moll, G., Black, M. J., SMPL: A Skinned Multi-Person Linear Model, ACM Transactions on Graphics, 2015. https://is.mpg.de/publications/smpl-2015",
        "[3] University of Wisconsin, Biovision BVH file format açıklaması. https://research.cs.wisc.edu/graphics/Courses/cs-838-1999/Jeff/BVH.html",
        "[4] Expo Documentation, React Native tabanlı mobil uygulama geliştirme dokümantasyonu. https://docs.expo.dev/",
        "[5] Fastify Documentation, Node.js için düşük overhead backend framework. https://fastify.dev/",
        "[6] PostgreSQL Documentation, JSON/JSONB veri tipleri ve indeksleme. https://www.postgresql.org/docs/16/datatype-json.html",
        "[7] Amazon S3 Documentation, presigned URL ile nesne indirme/yükleme. https://docs.aws.amazon.com/AmazonS3/latest/userguide/using-presigned-url.html",
        "[8] MinIO Documentation, S3 uyumlu object storage yaklaşımı. https://min.io/product/s3-compatibility",
        "[9] RunPod Documentation, serverless GPU iş yükleri ve model inference dağıtımı. https://docs.runpod.io/serverless/quick-deploys",
        "[10] ISO/IEC/IEEE 29148, Requirements engineering süreçleri ve gereksinim belirtimi standardı. https://www.iso.org/standard/72089.html",
        "[11] ISO/IEC/IEEE 12207, Yazılım yaşam döngüsü süreçleri standardı. https://www.iso.org/standard/63712.html",
        "[12] Semantic Versioning 2.0.0, yazılım sürümleme ilkeleri. https://semver.org/",
        "[13] OWASP Application Security Verification Standard, uygulama güvenliği kontrol yaklaşımı. https://owasp.org/www-project-application-security-verification-standard/",
        "[14] OWASP API Security Top 10, API güvenliği riskleri. https://owasp.org/API-Security/editions/2023/en/0x00-header/",
        "[15] Kişisel Verileri Koruma Kurumu, KVKK ve kişisel veri güvenliği rehberleri. https://www.kvkk.gov.tr/",
        "[16] Strategyzer, Business Model Canvas yaklaşımı. https://www.strategyzer.com/library/the-business-model-canvas",
        "[17] Ries, E., The Lean Startup yaklaşımı ve MVP kavramı. http://theleanstartup.com/",
    ]
    for ref in references:
        p = document.add_paragraph(style=body_style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    document.add_page_break()

    add_heading(document, "EKLER", front_h_style)
    add_heading(document, "EK A: Karar kuralları", front_h_style)
    for text in [
        "Bu ek, raporun son haline getirilirken hangi görsel, tablo, test kanıtı ve teknik rapor çıktılarının eklenmesi gerektiğini ayrıntılandırır. Amaç, kitapçığın yalnızca açıklayıcı metinlerden oluşmaması; projenin gerçekten çalıştığını, hangi seviyede hazır olduğunu ve hangi alanlarda geliştirme gerektiğini kanıtlayan teknik içeriklerle desteklenmesidir.",
        "Raporun ana gövdesinde mimari diyagramlar, veri akışları, deney düzeneği görselleri, job state akışı ve güvenlik kontrol akışı bulunmalıdır. Tablolarda ise gereksinim listesi, teknoloji seçimi, veri modeli, test senaryoları, kalite metrikleri, risk matrisi, maliyet/iş-zaman çizelgesi ve çıktı artifact listesi verilmelidir. Ekran görüntüleri koyulurken signed URL, token, bucket credential, veritabanı URL'si ve özel dosya yolları maskelenmelidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_caption(document, "Tablo A.1. Rapora eklenmesi önerilen kanıt ve çıktı listesi", caption_style)
    add_table(
        document,
        ["Eklenecek içerik", "Nereye eklenmeli", "Detay"],
        [
            ["Uygulama ekran görüntüleri", "Bölüm 1 ve 3", "Capture, Motion Preview, Upload Progress, Processing Status, Export Result ekranları."],
            ["Sistem mimarisi diyagramı", "Bölüm 2", "Mobil, backend API, PostgreSQL, S3/MinIO, worker ve RunPod ilişkisi."],
            ["Veri modeli tablosu", "Bölüm 2", "projects, takes, capture_sessions, capture_videos, processing_jobs, export_files."],
            ["API endpoint özeti", "Bölüm 2 veya Ek A", "Project/take/session/upload/process/export endpoint'leri ve görevleri."],
            ["Capture metadata örneği", "Bölüm 2 ve 4", "schema, deviceIndex, captureMode, video, camera, quality, sync alanları."],
            ["Quality report özeti", "Bölüm 3 ve 5", "score, grade, warnings, errors, validation, multiView alanları."],
            ["Motion pipeline report", "Bölüm 3 ve 5", "stage listesi, finalAnimationSource, whamInputUsage, artifact refs."],
            ["Blender import kanıtı", "Bölüm 3", "BVH dosyasının Blender'da açıldığını gösteren ekran görüntüsü veya smoke test sonucu."],
            ["İş-zaman çizelgesi", "Ek A", "Sprintler, görev sahipleri, teslim tarihleri ve gerçekleşme durumu."],
            ["Maliyet analizi", "Ek A", "Mobil cihaz, backend, storage, GPU/RunPod, geliştirme ve test maliyet kalemleri."],
            ["Risk/önlem matrisi", "Bölüm 4", "Güvenlik, veri bütünlüğü, ağ, model runtime ve lisans riskleri."],
        ],
        [1.7, 1.7, 3.5],
    )
    for text in [
        "Raporun tablo kısmına özellikle test senaryoları eklenmelidir. Her test satırı için test adı, ön koşul, giriş verisi, beklenen çıktı, gözlenen çıktı, başarı durumu ve kanıt dosyası sütunları kullanılmalıdır. Bu format, sözlü sınavda projenin hangi özelliklerinin gerçekten doğrulandığını açıkça göstermeyi sağlar.",
        "Kalite raporları için ayrı bir özet tablo hazırlanmalıdır. Bu tabloda export dosya adı, artifact formatı, dosya amacı, oluşma koşulu, kaliteye katkısı ve kullanıcı arayüzünde nerede gösterildiği yazılmalıdır. Örneğin bvh dosyası animasyon aktarımı için, smpl_parameters_json model çıktısını saklamak için, quality_report_json teknik kaliteyi göstermek için, motion_pipeline_report_json ise işlem hattının izlenebilirliği için gereklidir.",
        "Maliyet ve iş-zaman çizelgesi de rapora eklenmelidir. Maliyet tablosunda mevcut geliştirme cihazları, backend sunucu, object storage, GPU worker, test cihazları ve lisanslı model asset'leri ayrı kalemler halinde verilmelidir. İş-zaman çizelgesinde analiz, mobil kayıt, backend upload, worker pipeline, export, QA, raporlama ve sunum hazırlığı aşamaları gösterilmelidir.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    add_appendix_expansion(document, body_style, caption_style)
    document.add_page_break()

    add_heading(document, "ÖZGEÇMİŞ", front_h_style)
    for text in [
        "Burak Coşkun, Sakarya Üniversitesi Bilgisayar ve Bilişim Bilimleri Fakültesi Bilgisayar Mühendisliği Bölümü öğrencisidir. Mobil uygulama geliştirme, backend servisleri, yapay zeka destekli uygulamalar ve görüntü işleme tabanlı sistemler üzerine çalışmalar yürütmektedir. Bu bitirme çalışmasında mobil uygulama mimarisi, backend entegrasyonu, upload ve processing akışlarının tasarlanması konularında görev almıştır.",
        "Berke Pite, Sakarya Üniversitesi Bilgisayar ve Bilişim Bilimleri Fakültesi Bilgisayar Mühendisliği Bölümü öğrencisidir. Projede mobil kullanıcı deneyimi, test senaryoları, dokümantasyon, çoklu kamera akışlarının değerlendirilmesi ve raporlama süreçlerinde görev almıştır.",
        "Eren Ozan Pite, Sakarya Üniversitesi Bilgisayar ve Bilişim Bilimleri Fakültesi Bilgisayar Mühendisliği Bölümü öğrencisidir. Projede backend iş akışı, veri modeli, kalite raporları, test planı ve sistem değerlendirme süreçlerinde görev almıştır.",
    ]:
        add_paragraph(document, text, body_style, first=True)
    document.add_page_break()

    add_heading(document, "BSM 498 BİTİRME ÇALIŞMASI DEĞERLENDİRME VE SÖZLÜ SINAV TUTANAĞI", front_h_style)
    front_table(
        document,
        [
            ("KONU", "Mobil cihazlar ile video tabanlı işaretleyicisiz hareket yakalama sistemi geliştirmesi"),
            ("ÖĞRENCİLER", "B231210351 - Burak Coşkun; Öğrenci2 No - Berke Pite; B221210059 - Eren Ozan Pite"),
            ("DANIŞMAN", "Danışman adı eklenecektir"),
            ("NOT", "Bu sayfa jüri değerlendirmesi için şablonda bırakılmıştır."),
        ],
    )

    for style in document.styles:
        if style.type == 1:
            try:
                style.font.name = "Times New Roman"
            except Exception:
                pass

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUT))


if __name__ == "__main__":
    build_report()
    print(OUT)
